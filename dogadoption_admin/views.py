"""Administrative views for the dog adoption dashboard.

The file is documented and separated by admin navigation groups so related
features are easier to find and maintain.
"""

from collections import defaultdict
from datetime import datetime, time, timedelta
from decimal import Decimal
from functools import wraps
import hashlib
import io
import json
import re
import secrets
from types import SimpleNamespace
from urllib.parse import urlencode

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
except ImportError:
    Workbook = Alignment = Border = Font = PatternFill = Side = None

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import landscape, letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
except ImportError:
    colors = landscape = letter = getSampleStyleSheet = Paragraph = None
    SimpleDocTemplate = Spacer = Table = TableStyle = None

from django.contrib import messages
from django.contrib.auth import logout
from django.contrib.auth import update_session_auth_hash
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.core.cache import cache
from django.core.paginator import Paginator
from django.db import DatabaseError, transaction
from django.db.models import Case, CharField, Count, DateField, F, IntegerField, Max, Min, OuterRef, Prefetch, Q, Subquery, Value, When
from django.db.models.functions import Cast, Coalesce, Concat, Lower, Substr, Trim, TruncDate
from django.http import Http404, HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.template.loader import render_to_string
from django.templatetags.static import static
from django.urls import reverse
from django.utils.http import url_has_allowed_host_and_scheme
from django.utils import timezone
from django.utils.dateparse import parse_date, parse_datetime
from django.views.decorators.http import require_http_methods, require_POST

from .access import (
    STAFF_PERMISSION_FIELDS,
    STAFF_PERMISSION_GROUPS,
    get_admin_access,
    get_staff_landing_url,
    get_staff_permission_summary,
    is_route_allowed,
)
from .forms import CitationForm, ManagedStaffAccountForm, PenaltyForm, PostForm, SectionForm
from .admin_notification_utils import sync_expiry_notifications
from .barangays import BAYAWAN_BARANGAYS
from .cache_utils import ANALYTICS_DASHBOARD_CACHE_KEY
from .context_processors import (
    ADMIN_NOTIFICATIONS_CACHE_KEY,
    ADMIN_NOTIFICATIONS_CACHE_TTL_SECONDS,
)
from .models import (
    AdminNotification,
    Barangay,
    CertificateSettings,
    Citation,
    DewormingTreatmentRecord,
    Dog,
    DogImage,
    DogAnnouncement,
    DogAnnouncementImage,
    DogRegistration,
    GlobalAppointmentDate,
    Penalty,
    PenaltySection,
    Post,
    PostImage,
    PostRequest,
    StaffAccess,
    UserViolationNotification,
    UserViolationSummary,
    VaccinationRecord,
)
from user.models import (
    DogCaptureRequest,
    DogCaptureRequestImage,
    MissingDogPost,
    Profile,
    UserAdoptionImage,
    UserAdoptionPost,
)


def _build_admin_notification_summary():
    cached = cache.get(ADMIN_NOTIFICATIONS_CACHE_KEY)
    if cached is not None:
        return cached

    payload = {
        "admin_pending_capture_count": DogCaptureRequest.objects.filter(status="pending").count(),
        "admin_unread_notifications": AdminNotification.objects.filter(is_read=False).count(),
        "admin_latest_notifications": list(
            AdminNotification.objects.order_by("-created_at")
            .values("id", "title", "message", "created_at", "is_read")[:5]
        ),
    }
    cache.set(
        ADMIN_NOTIFICATIONS_CACHE_KEY,
        payload,
        ADMIN_NOTIFICATIONS_CACHE_TTL_SECONDS,
    )
    return payload
from user.notification_utils import (
    bump_user_home_feed_namespace,
    invalidate_user_notification_content,
    invalidate_user_notification_payload,
    remember_request_reviewed_at,
)


POST_HISTORY_CACHE_KEY = "dogadoption_admin_post_history_ids_v1"
POST_HISTORY_CACHE_TTL_SECONDS = 120
HOME_SPOTLIGHT_PIN_LIMIT = 4
VIOLATION_WARNING_THRESHOLD = 3
VIOLATION_OFFICE_NAME = "CITY VETERINARY OFFICE"
VIOLATION_OFFICE_ADDRESS_LINES = (
    "National Highway, Barangay Villareal, Bayawan City",
    "Negros Oriental, 6221 Philippines",
)
VIOLATION_SIGNATORY_NAME = "REYNALDO SOLAMILLO"
VIOLATION_SIGNATORY_ROLE = "Team Leader-Rabies Control Team"
ADMIN_USERS_PER_PAGE = 25


# =============================================================================
# Shared imports, constants, and helper utilities
# =============================================================================


def _get_cached_post_history_ids():
    history_candidate_ids = cache.get(POST_HISTORY_CACHE_KEY)
    if history_candidate_ids is not None:
        return history_candidate_ids

    finalized_ids = list(
        Post.objects.filter(status__in=["reunited", "adopted"])
        .order_by("-created_at", "-id")
        .values_list("id", flat=True)
    )
    archived_ids = list(
        Post.objects.filter(is_history=True)
        .order_by("-created_at", "-id")
        .values_list("id", flat=True)
    )
    history_candidate_posts = list(
        Post.with_pending_request_state(
            Post.objects.filter(
                is_history=False,
                status__in=["rescued", "under_care"],
            )
        )
        .only(
            "id",
            "status",
            "created_at",
            "claim_days",
            "is_history",
            "phase_override",
            "phase_override_started_at",
        )
        .order_by("-created_at")
    )
    Post.attach_active_appointment_dates(history_candidate_posts)
    expired_ids = [
        post.id
        for post in history_candidate_posts
        if post.is_expired()
    ]
    seen_ids = set()
    history_candidate_ids = []
    for post_id in [*finalized_ids, *archived_ids, *expired_ids]:
        if post_id in seen_ids:
            continue
        seen_ids.add(post_id)
        history_candidate_ids.append(post_id)
    cache.set(
        POST_HISTORY_CACHE_KEY,
        history_candidate_ids,
        POST_HISTORY_CACHE_TTL_SECONDS,
    )
    return history_candidate_ids


def _build_post_history_page(request, page_param="page", rows_per_page=10):
    history_candidate_ids = _get_cached_post_history_ids()
    filter_type = (request.GET.get("record_type") or "all").strip().lower()
    if filter_type not in {"all", "adopted", "redeemed", "unresolved"}:
        filter_type = "all"
    history_return_to = request.get_full_path()

    history_meta_map = {
        post.id: post
        for post in Post.objects.filter(id__in=history_candidate_ids).only("id", "status", "is_history")
    }

    adopted_total = 0
    redeemed_total = 0
    unresolved_total = 0
    for post_id in history_candidate_ids:
        post = history_meta_map.get(post_id)
        if not post:
            continue
        if post.status == "adopted":
            adopted_total += 1
        elif post.status == "reunited":
            redeemed_total += 1
        else:
            unresolved_total += 1

    history_total = len(history_candidate_ids)

    def _matches_filter(post):
        if filter_type == "adopted":
            return post.status == "adopted"
        if filter_type == "redeemed":
            return post.status == "reunited"
        if filter_type == "unresolved":
            return post.status not in {"adopted", "reunited"}
        return True

    filtered_history_ids = [
        post_id
        for post_id in history_candidate_ids
        if (post := history_meta_map.get(post_id)) and _matches_filter(post)
    ]

    paginator = Paginator(filtered_history_ids, rows_per_page)
    page_obj = paginator.get_page(request.GET.get(page_param, 1))
    page_ids = list(page_obj.object_list)
    history_posts = []

    if page_ids:
        post_map = {
            post.id: post
            for post in Post.objects.filter(id__in=page_ids)
            .only(
                "id",
                "caption",
                "breed",
                "breed_other",
                "age_group",
                "size_group",
                "gender",
                "coat_length",
                "colors",
                "color_other",
                "location",
                "status",
                "created_at",
                "claim_days",
                "is_history",
                "phase_override",
                "phase_override_started_at",
            )
        }
        Post.attach_active_appointment_dates(post_map.values())
        primary_image_by_post_id = {}
        for image in PostImage.objects.filter(post_id__in=page_ids).only("post_id", "image").order_by("id"):
            if image.post_id in primary_image_by_post_id:
                continue
            image_url = _safe_media_url(image.image)
            if image_url:
                primary_image_by_post_id[image.post_id] = image_url

        accepted_request_map = {}
        accepted_requests = list(
            PostRequest.objects.filter(
                post_id__in=page_ids,
                status="accepted",
                request_type__in=["claim", "adopt"],
            )
            .select_related("user")
            .only(
                "id",
                "post_id",
                "request_type",
                "created_at",
                "scheduled_appointment_date",
                "user__username",
                "user__first_name",
                "user__last_name",
            )
            .order_by("-created_at", "-id")
        )
        for req in accepted_requests:
            request_bucket = accepted_request_map.setdefault(req.post_id, {})
            request_bucket.setdefault(req.request_type, req)

        for post_id in page_ids:
            post = post_map.get(post_id)
            if not post:
                continue
            is_archived = bool(getattr(post, "is_history", False))
            accepted_req = None
            record_label = "Not Adopted / Redeemed"
            record_tone = "warning"
            record_source_label = "Expired listing"
            record_note = "No final adoption or redeem record."
            recorded_on = post.adoption_deadline()

            if post.status == "adopted":
                accepted_req = accepted_request_map.get(post_id, {}).get("adopt")
                adopter_name = "-"
                if accepted_req:
                    adopter_name = (
                        f"{(accepted_req.user.first_name or '').strip()} {(accepted_req.user.last_name or '').strip()}".strip()
                        or accepted_req.user.username
                    )
                record_label = "Adopted"
                record_tone = "adopted"
                if accepted_req:
                    record_source_label = "Finalized adoption"
                    record_note = f"Adopted by {adopter_name}"
                else:
                    record_source_label = "Admin-completed adoption"
                    record_note = "Marked as adopted by admin without a request record."
                recorded_on = (
                    accepted_req.scheduled_appointment_date
                    if accepted_req and accepted_req.scheduled_appointment_date
                    else accepted_req.created_at if accepted_req else post.created_at
                )
            elif post.status == "reunited":
                accepted_req = accepted_request_map.get(post_id, {}).get("claim")
                owner_name = "-"
                if accepted_req:
                    owner_name = (
                        f"{(accepted_req.user.first_name or '').strip()} {(accepted_req.user.last_name or '').strip()}".strip()
                        or accepted_req.user.username
                    )
                record_label = "Redeemed"
                record_tone = "reunited"
                if accepted_req:
                    record_source_label = "Finalized redeem"
                    record_note = f"Redeemed by {owner_name}"
                else:
                    record_source_label = "Admin-completed redeem"
                    record_note = "Marked as redeemed by admin without a request record."
                recorded_on = (
                    accepted_req.scheduled_appointment_date
                    if accepted_req and accepted_req.scheduled_appointment_date
                    else accepted_req.created_at if accepted_req else post.created_at
                )
            elif is_archived:
                record_source_label = "Archived record"
                record_note = "Moved to history without a final adoption or redeem."
                recorded_on = post.adoption_deadline() or post.created_at

            detail_rows = []
            detail_sources = [
                ("Breed", post.display_breed),
                ("Color", post.display_colors),
                ("Age", post.display_age_group),
                ("Size", post.display_size_group),
                ("Coat", post.display_coat_length),
                ("Gender", post.get_gender_display() if post.gender else ""),
                ("Location", post.location or ""),
            ]
            for label, value in detail_sources:
                cleaned_value = (value or "").strip()
                if cleaned_value:
                    detail_rows.append({
                        "label": label,
                        "value": cleaned_value,
                    })

            record_actions = []
            if record_tone == "warning":
                record_actions = [
                    {
                        "status": "adopted",
                        "label": "Record Adopted",
                        "confirm_title": "Record as Adopted",
                        "confirm_message": "Record this expired post as adopted and keep it in history?",
                        "confirm_submit_label": "Record Adopted",
                        "button_class": "history-action-btn--adopted",
                        "icon": "fas fa-heart",
                    },
                    {
                        "status": "reunited",
                        "label": "Record Redeemed",
                        "confirm_title": "Record as Redeemed",
                        "confirm_message": "Record this expired post as redeemed and keep it in history?",
                        "confirm_submit_label": "Record Redeemed",
                        "button_class": "history-action-btn--redeemed",
                        "icon": "fas fa-hand-holding-heart",
                    },
                ]

            history_posts.append({
                "post": post,
                "record_label": record_label,
                "record_tone": record_tone,
                "record_source_label": record_source_label,
                "record_note": record_note,
                "base_status_label": post.get_status_display(),
                "recorded_on": recorded_on,
                "record_date_label": "Expired on" if record_tone == "warning" else "Recorded on",
                "primary_image_url": primary_image_by_post_id.get(post_id, ""),
                "post_detail_rows": detail_rows,
                "post_title": post.display_title,
                "can_record_final_status": bool(record_actions),
                "record_actions": record_actions,
                "record_action_url": reverse("dogadoption_admin:record_history_status", args=[post.id]),
            })

    return {
        "history_total": history_total,
        "history_adopted_total": adopted_total,
        "history_redeemed_total": redeemed_total,
        "history_unresolved_total": unresolved_total,
        "history_active_filter": filter_type,
        "history_filtered_total": len(filtered_history_ids),
        "history_posts": history_posts,
        "history_page_obj": page_obj,
        "history_return_to": history_return_to,
    }

CAT_BREED_KEYWORDS = {
    "abyssinian",
    "american curl",
    "american shorthair",
    "balinese",
    "bengal",
    "birman",
    "bombay",
    "british shorthair",
    "burmese",
    "chartreux",
    "cornish rex",
    "devon rex",
    "domestic longhair",
    "domestic shorthair",
    "egyptian mau",
    "exotic shorthair",
    "feline",
    "himalayan",
    "maine coon",
    "manx",
    "munchkin",
    "norwegian forest",
    "ocicat",
    "oriental shorthair",
    "persian",
    "puspin",
    "ragdoll",
    "russian blue",
    "savannah",
    "scottish fold",
    "selkirk rex",
    "siamese",
    "siberian",
    "singapura",
    "snowshoe",
    "sphynx",
    "tonkinese",
    "turkish angora",
    "turkish van",
}

ACTIVE_BARANGAY_LOOKUP_CACHE_KEY = "dogadoption_admin_active_barangay_lookup"
ACTIVE_BARANGAY_LOOKUP_CACHE_TTL_SECONDS = 300


def _get_python_docx_document():
    """Return the Word export helper or raise a clear dependency error."""
    if Document is None:
        raise RuntimeError("python-docx is required for Word export.")
    return Document


def _get_openpyxl_exports():
    """Return Excel export helpers or raise a clear dependency error."""
    if Workbook is None:
        raise RuntimeError("openpyxl is required for Excel export.")
    return Workbook, Alignment, Border, Font, PatternFill, Side


def _get_reportlab_exports():
    """Return PDF export helpers or raise a clear dependency error."""
    if SimpleDocTemplate is None:
        raise RuntimeError("reportlab is required for PDF export.")
    return colors, landscape, letter, getSampleStyleSheet, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


def _clean_barangay(value):
    return " ".join((value or "").split()).strip()


def _normalize_person_name(value):
    return " ".join((value or "").split()).strip().casefold()


def _clean_breed(value):
    return " ".join((value or "").split()).strip()


def _normalize_breed_key(value):
    return re.sub(r"[^a-z0-9]+", " ", _clean_breed(value).casefold()).strip()


def _format_breed_label(value):
    cleaned = _clean_breed(value)
    if not cleaned:
        return ""
    if cleaned == cleaned.lower() or cleaned == cleaned.upper():
        return cleaned.title()
    return cleaned


def _normalize_certificate_series(value):
    parts = [
        re.sub(r"[^A-Za-z0-9]+", "", part).upper()
        for part in re.split(r"[-/\s]+", (value or "").strip())
    ]
    parts = [part for part in parts if part]
    if not parts:
        return ""

    if parts[0] != "CVET":
        parts.insert(0, "CVET")

    if len(parts) >= 3 and parts[-1].isdigit():
        parts = parts[:-1]

    return "-".join(parts)


def _next_certificate_sequence(series_prefix):
    """Return the next numeric suffix for a certificate series without scanning every row."""
    pattern = re.compile(rf"^{re.escape(series_prefix)}-(\d+)$")
    prefix_len = len(series_prefix)
    start_pos = prefix_len + 2  # first digit after "{series_prefix}-"
    qs = DogRegistration.objects.filter(reg_no__startswith=f"{series_prefix}-").filter(
        reg_no__regex=rf"^{re.escape(series_prefix)}-[0-9]+$",
    )
    try:
        max_sequence = qs.annotate(
            _seq=Cast(Substr("reg_no", start_pos), output_field=IntegerField()),
        ).aggregate(max_seq=Max("_seq"))["max_seq"]
    except (DatabaseError, ValueError, TypeError):
        max_sequence = None
    if max_sequence is not None:
        return int(max_sequence) + 1

    max_sequence = 0
    for reg_no in DogRegistration.objects.filter(
        reg_no__startswith=f"{series_prefix}-"
    ).values_list("reg_no", flat=True):
        match = pattern.match((reg_no or "").upper())
        if match:
            max_sequence = max(max_sequence, int(match.group(1)))

    return max_sequence + 1


def _build_certificate_registration_number(series_prefix):
    next_sequence = _next_certificate_sequence(series_prefix)
    return f"{series_prefix}-{next_sequence}"


def _exclude_breed_from_chart(value):
    return _normalize_breed_key(value) == "mongril"


def _classify_breed_type(value):
    breed_key = _normalize_breed_key(value)
    if not breed_key:
        return "dog"
    if any(keyword in breed_key for keyword in CAT_BREED_KEYWORDS):
        return "cat"
    return "dog"


def _owner_initials(name):
    parts = [part for part in (name or "").strip().split() if part]
    if not parts:
        return "?"
    if len(parts) == 1:
        return parts[0][:1].upper()
    return f"{parts[0][:1]}{parts[-1][:1]}".upper()


def _get_active_barangay_lookup():
    lookup = cache.get(ACTIVE_BARANGAY_LOOKUP_CACHE_KEY)
    if lookup is None:
        lookup = {
            _normalize_barangay(name): name
            for name in Barangay.objects.filter(is_active=True).values_list("name", flat=True)
        }
        cache.set(
            ACTIVE_BARANGAY_LOOKUP_CACHE_KEY,
            lookup,
            ACTIVE_BARANGAY_LOOKUP_CACHE_TTL_SECONDS,
        )
    return lookup


def _safe_media_url(file_field):
    if not file_field:
        return ""
    try:
        return file_field.url
    except (AttributeError, ValueError):
        return ""


def _dog_image_prefetch():
    return Prefetch(
        "images",
        queryset=DogImage.objects.only("id", "dog_id", "image").order_by("created_at", "id"),
    )


def _registered_dog_payload(dog):
    photo_urls = []
    for image in dog.images.all():
        image_url = _safe_media_url(getattr(image, "image", None))
        if image_url:
            photo_urls.append(image_url)

    return {
        "id": dog.id,
        "name": dog.name or "Unnamed Dog",
        "species": dog.species or "Canine",
        "sex_label": dog.get_sex_display() if dog.sex else "-",
        "age": dog.age or "-",
        "neutering_label": dog.get_neutering_status_display() if dog.neutering_status else "-",
        "color": dog.color or "-",
        "date_registered": dog.date_registered,
        "location": dog.barangay or dog.owner_address or "",
        "photo_urls": photo_urls,
    }


def _build_registered_dog_payloads(dogs):
    return [_registered_dog_payload(dog) for dog in dogs]


def _build_owner_profile_lookup(owner_names):
    normalized_names = {_normalize_person_name(name) for name in owner_names if name}
    if not normalized_names:
        return {}

    profiles = (
        Profile.objects.select_related("user")
        .annotate(
            owner_full_name_norm=Lower(
                Trim(
                    Concat(
                        "user__first_name",
                        Value(" "),
                        "user__last_name",
                    )
                )
            )
        )
        .filter(
            owner_full_name_norm__in=normalized_names,
            user__is_active=True,
            user__is_staff=False,
        )
        .only("user_id", "profile_image", "user__first_name", "user__last_name")
    )

    grouped_profiles = defaultdict(list)
    for profile in profiles:
        grouped_profiles[profile.owner_full_name_norm].append(profile)

    lookup = {}
    for normalized_name, matches in grouped_profiles.items():
        # Duplicate-name user accounts are ambiguous, so do not attach
        # a manual registration row to any specific profile in that case.
        if len(matches) != 1:
            continue

        profile = matches[0]
        image_url = _safe_media_url(getattr(profile, "profile_image", None))
        lookup[normalized_name] = {
            "image_url": image_url,
            "user_id": profile.user_id,
        }
    return lookup


def _normalize_barangay(value):
    return "".join(ch.lower() for ch in _clean_barangay(value) if ch.isalnum())


def _resolve_barangay_name(value):
    normalized = _normalize_barangay(value)
    if not normalized:
        return ""
    return _get_active_barangay_lookup().get(normalized, "")


#extracting barangays
def _extract_barangay_from_address(address):
    cleaned = _clean_barangay(address)
    if not cleaned:
        return ""

    parts = [p.strip() for p in cleaned.split(",") if p.strip()]
    if len(parts) >= 3 and parts[-2].lower() == "bayawan city" and parts[-1].lower() == "negros oriental":
        candidate = parts[-3]
        resolved = _resolve_barangay_name(candidate)
        return resolved or candidate

    for part in reversed(parts):
        resolved = _resolve_barangay_name(part)
        if resolved:
            return resolved

    return _resolve_barangay_name(cleaned)


BAYAWAN_ALLOWED_BARANGAYS = BAYAWAN_BARANGAYS

BAYAWAN_ALLOWED_BARANGAY_KEYS = {
    _normalize_barangay(name) for name in BAYAWAN_ALLOWED_BARANGAYS
}

REGISTRATION_RECORD_LOCATOR_POINTS = {
    "malabugas": {"x": 37.0, "y": 12.0},
    "maninihon": {"x": 51.0, "y": 18.0},
    "nangka": {"x": 69.0, "y": 12.0},
    "kalumboyan": {"x": 63.0, "y": 23.0},
    "banaybanay": {"x": 35.0, "y": 24.0},
    "aliis": {"x": 23.0, "y": 29.0},
    "banga": {"x": 42.0, "y": 34.0},
    "pagatban": {"x": 59.0, "y": 34.0},
    "bugay": {"x": 75.0, "y": 31.0},
    "boyco": {"x": 29.0, "y": 43.0},
    "minaba": {"x": 43.0, "y": 45.0},
    "narra": {"x": 57.0, "y": 46.0},
    "cansumalig": {"x": 70.0, "y": 45.0},
    "dawis": {"x": 81.0, "y": 42.0},
    "kalamtukan": {"x": 35.0, "y": 55.0},
    "tayawan": {"x": 49.0, "y": 57.0},
    "villareal": {"x": 62.0, "y": 58.0},
    "poblacion": {"x": 74.0, "y": 60.0},
    "manduao": {"x": 29.0, "y": 70.0},
    "sanjose": {"x": 47.0, "y": 67.0},
    "sanisidro": {"x": 61.0, "y": 68.0},
    "suba": {"x": 78.0, "y": 69.0},
    "sanmiguel": {"x": 42.0, "y": 78.0},
    "sanroque": {"x": 56.0, "y": 80.0},
    "tinago": {"x": 70.0, "y": 77.0},
    "ubos": {"x": 82.0, "y": 80.0},
    "tabuan": {"x": 37.0, "y": 90.0},
    "villasol": {"x": 54.0, "y": 91.0},
}


def _normalize_city(value):
    return "".join(ch.lower() for ch in _clean_barangay(value) if ch.isalnum())


def _build_registration_locator_points(barangay_names):
    """Create schematic locator positions for registration record barangay highlighting."""
    points = []
    fallback_index = 0

    for name in barangay_names:
        normalized_name = _normalize_barangay(name)
        coords = REGISTRATION_RECORD_LOCATOR_POINTS.get(normalized_name)

        if coords is None:
            col = fallback_index % 4
            row = fallback_index // 4
            coords = {
                "x": 18.0 + (col * 18.0),
                "y": 18.0 + (row * 12.0),
            }
            fallback_index += 1

        points.append(
            {
                "name": name,
                "x": coords["x"],
                "y": coords["y"],
            }
        )

    return points


def _is_bayawan_city(value):
    return _normalize_city(value) in {"bayawan", "bayawancity"}


def _extract_city_from_address(address):
    cleaned = _clean_barangay(address)
    if not cleaned:
        return ""

    parts = [p.strip() for p in cleaned.split(",") if p.strip()]
    for part in reversed(parts):
        if _is_bayawan_city(part):
            return "Bayawan City"
    return ""


def _is_allowed_bayawan_map_point(barangay, city):
    return (
        _normalize_barangay(barangay) in BAYAWAN_ALLOWED_BARANGAY_KEYS
        and _is_bayawan_city(city)
    )


def _build_owner_full_name(first_name, last_name, fallback=""):
    first = " ".join((first_name or "").split()).strip()
    last = " ".join((last_name or "").split()).strip()
    fallback_clean = " ".join((fallback or "").split()).strip()

    if first or last:
        return f"{first} {last}".strip()
    return fallback_clean


def _registration_owner_key_from_names(first_name, last_name, fallback=""):
    owner_name = _build_owner_full_name(first_name, last_name, fallback)
    return _normalize_person_name(owner_name)


def _resolve_registration_owner_identity(owner_first_name, owner_last_name, owner_user_id=""):
    first = " ".join((owner_first_name or "").split()).strip()
    last = " ".join((owner_last_name or "").split()).strip()
    owner_name_key = _registration_owner_key_from_names(first, last)
    resolved_owner_user = None

    owner_user_id_text = str(owner_user_id or "").strip()
    if owner_user_id_text.isdigit() and owner_name_key:
        resolved_owner_user = (
            User.objects.filter(
                id=int(owner_user_id_text),
                is_active=True,
                is_staff=False,
                first_name__iexact=first,
                last_name__iexact=last,
            )
            .only("id", "first_name", "last_name")
            .first()
        )

    canonical_first = resolved_owner_user.first_name if resolved_owner_user else first
    canonical_last = resolved_owner_user.last_name if resolved_owner_user else last
    canonical_owner_name = _build_owner_full_name(canonical_first, canonical_last)
    canonical_owner_key = _normalize_person_name(canonical_owner_name)

    return canonical_owner_name, canonical_owner_key, resolved_owner_user


def _split_owner_name_parts(owner_name):
    cleaned = " ".join((owner_name or "").split()).strip()
    if not cleaned:
        return "", ""
    parts = cleaned.split()
    return parts[0], " ".join(parts[1:])


def _certificate_status_from_neutering(value):
    normalized = " ".join((value or "").split()).strip()
    return {
        "C": "Castrated",
        "S": "Spayed",
        "No": "Intact",
    }.get(normalized, "")


def _serialize_registered_pet(dog):
    pet_name = " ".join((getattr(dog, "name", "") or "").split()).strip()
    return {
        "id": getattr(dog, "id", None),
        "name": pet_name,
        "species": (getattr(dog, "species", "") or "").strip(),
        "sex": (getattr(dog, "sex", "") or "").strip(),
        "status": _certificate_status_from_neutering(getattr(dog, "neutering_status", "")),
        "color": (getattr(dog, "color", "") or "").strip(),
        "barangay": _resolve_barangay_name(getattr(dog, "barangay", "") or "") or "",
    }


def _owner_name_matches_query(owner_name, query_tokens, query):
    normalized_owner_name = _normalize_person_name(owner_name)
    normalized_query = _normalize_person_name(query)
    if not normalized_owner_name or not normalized_query:
        return False
    if normalized_query in normalized_owner_name:
        return True
    return all(
        _normalize_person_name(token) in normalized_owner_name
        for token in query_tokens
        if _normalize_person_name(token)
    )


def _build_registration_owner_search_results(query, limit):
    tokens = query.split()
    results = []

    users = User.objects.filter(is_active=True, is_staff=False)
    if len(tokens) >= 2:
        users = users.filter(
            (Q(first_name__istartswith=tokens[0]) & Q(last_name__istartswith=tokens[-1]))
            | Q(username__istartswith=query)
        )
    else:
        term = tokens[0]
        users = users.filter(
            Q(first_name__istartswith=term)
            | Q(last_name__istartswith=term)
            | Q(username__istartswith=term)
        )

    user_rows = list(
        users.order_by("first_name", "last_name", "id").values(
            "id",
            "first_name",
            "last_name",
            "username",
            "profile__address",
            "profile__phone_number",
        )[:limit]
    )

    user_ids = [row["id"] for row in user_rows]
    user_owner_keys = {}
    owner_key_to_user_id = {}
    for row in user_rows:
        owner_key = _registration_owner_key_from_names(
            row.get("first_name") or "",
            row.get("last_name") or "",
        )
        if owner_key:
            user_owner_keys[row["id"]] = owner_key
            owner_key_to_user_id.setdefault(owner_key, row["id"])

    registered_pets_by_user_id = defaultdict(list)
    seen_pet_keys_by_user_id = defaultdict(set)
    if user_ids:
        owner_keys = [key for key in user_owner_keys.values() if key]
        user_dogs = Dog.objects.filter(
            Q(owner_user_id__in=user_ids)
            | Q(owner_name_key__in=owner_keys)
        ).only(
            "id",
            "name",
            "species",
            "sex",
            "neutering_status",
            "color",
            "owner_user_id",
            "owner_name_key",
            "barangay",
        ).order_by("name", "id")

        for dog in user_dogs:
            owner_user_id = getattr(dog, "owner_user_id", None)
            if owner_user_id not in user_ids:
                owner_user_id = owner_key_to_user_id.get(
                    getattr(dog, "owner_name_key", "") or ""
                )
            if owner_user_id not in user_ids:
                continue

            pet_payload = _serialize_registered_pet(dog)
            pet_signature = (
                _normalize_person_name(pet_payload["name"]),
                pet_payload["species"],
                pet_payload["sex"],
                pet_payload["status"],
                _normalize_person_name(pet_payload["color"]),
            )
            if pet_signature in seen_pet_keys_by_user_id[owner_user_id]:
                continue
            seen_pet_keys_by_user_id[owner_user_id].add(pet_signature)
            registered_pets_by_user_id[owner_user_id].append(pet_payload)

    for row in user_rows:
        first_name = (row.get("first_name") or "").strip()
        last_name = (row.get("last_name") or "").strip()
        username = (row.get("username") or "").strip()
        full_name = f"{first_name} {last_name}".strip()
        barangay = _extract_barangay_from_address(row.get("profile__address") or "")
        phone_number = (row.get("profile__phone_number") or "").strip()
        registered_pets = registered_pets_by_user_id.get(row["id"], [])
        results.append(
            {
                "id": row["id"],
                "owner_key": f"user:{row['id']}",
                "source": "user",
                "source_label": "User account",
                "first_name": first_name,
                "last_name": last_name,
                "username": username,
                "full_name": full_name or username,
                "barangay": barangay,
                "phone_number": phone_number,
                "registered_pet_count": len(registered_pets),
                "registered_pets": registered_pets,
            }
        )

    remaining_slots = max(limit - len(results), 0)
    if remaining_slots:
        manual_owner_filter = Q(owner_name__istartswith=query)
        normalized_query = _normalize_person_name(query)
        if normalized_query:
            manual_owner_filter |= Q(owner_name_key__istartswith=normalized_query)
        for token in tokens:
            normalized_token = _normalize_person_name(token)
            manual_owner_filter |= Q(owner_name__icontains=token)
            if normalized_token:
                manual_owner_filter |= Q(owner_name_key__icontains=normalized_token)

        manual_dogs = list(
            Dog.objects.filter(owner_user__isnull=True)
            .filter(manual_owner_filter)
            .only(
                "id",
                "name",
                "species",
                "sex",
                "neutering_status",
                "color",
                "owner_name",
                "owner_name_key",
                "owner_address",
                "barangay",
            )
            .order_by("owner_name", "name", "id")[: max(remaining_slots * 8, 24)]
        )

        seen_owner_keys = {user_owner_keys.get(row["id"], "") for row in user_rows}
        manual_results_map = {}
        manual_pet_keys = defaultdict(set)
        for dog in manual_dogs:
            owner_name = " ".join((getattr(dog, "owner_name", "") or "").split()).strip()
            owner_name_key = getattr(dog, "owner_name_key", "") or _normalize_person_name(owner_name)
            if not owner_name_key or owner_name_key in seen_owner_keys:
                continue
            if not _owner_name_matches_query(owner_name, tokens, query):
                continue

            pet_payload = _serialize_registered_pet(dog)
            pet_signature = (
                _normalize_person_name(pet_payload["name"]),
                pet_payload["species"],
                pet_payload["sex"],
                pet_payload["status"],
                _normalize_person_name(pet_payload["color"]),
            )
            if pet_signature in manual_pet_keys[owner_name_key]:
                continue
            manual_pet_keys[owner_name_key].add(pet_signature)

            result_row = manual_results_map.setdefault(
                owner_name_key,
                {
                    "id": None,
                    "owner_key": f"name:{owner_name_key}",
                    "source": "manual_owner",
                    "source_label": "Registered owner",
                    "first_name": _split_owner_name_parts(owner_name)[0],
                    "last_name": _split_owner_name_parts(owner_name)[1],
                    "username": "",
                    "full_name": owner_name,
                    "barangay": _resolve_barangay_name(getattr(dog, "barangay", "") or "")
                    or _extract_barangay_from_address(getattr(dog, "owner_address", "") or ""),
                    "phone_number": "",
                    "registered_pet_count": 0,
                    "registered_pets": [],
                },
            )
            result_row["registered_pets"].append(pet_payload)
            result_row["registered_pet_count"] = len(result_row["registered_pets"])

        for owner_result in manual_results_map.values():
            results.append(owner_result)
            if len(results) >= limit:
                break

    return results[:limit]


def _build_owner_limit_query(owner_name_key, owner_name, owner_user=None):
    normalized_owner_key = _normalize_person_name(owner_name_key or owner_name)
    owner_query = Q()

    if normalized_owner_key:
        owner_query |= Q(owner_name_key=normalized_owner_key)

    if owner_name:
        owner_query |= Q(owner_name__iexact=owner_name)

    if owner_user is not None:
        owner_query |= Q(owner_user=owner_user)

    return owner_query


def _build_registration_record_owner_key(dog, matched_owner_user_id=None):
    owner_user_id = getattr(dog, "owner_user_id", None) or matched_owner_user_id
    if owner_user_id:
        return f"user:{owner_user_id}"

    normalized_owner = _normalize_person_name(
        getattr(dog, "owner_name_key", "") or getattr(dog, "owner_name", "")
    )
    if normalized_owner:
        return f"name:{normalized_owner}"

    return f"dog:{getattr(dog, 'id', 'unknown')}"


def _certificate_reg_no_sort_key(reg_no):
    reg_no_text = " ".join(str(reg_no or "").split()).strip()
    numeric_parts = [int(part) for part in re.findall(r"\d+", reg_no_text)]
    if not numeric_parts:
        return (0, 0, 0, reg_no_text.casefold())

    numeric_parts = numeric_parts[-3:]
    while len(numeric_parts) < 3:
        numeric_parts.insert(0, 0)
    return (*numeric_parts, reg_no_text.casefold())


def _build_certificate_manual_owner_preview_keys(owner_names):
    normalized_names = {_normalize_person_name(name) for name in owner_names if name}
    if not normalized_names:
        return set()

    return set(
        Dog.objects.annotate(
            owner_lookup_key=Case(
                When(
                    Q(owner_name_key__isnull=False) & ~Q(owner_name_key=""),
                    then=Lower(Trim("owner_name_key")),
                ),
                default=Lower(Trim("owner_name")),
                output_field=CharField(),
            )
        )
        .filter(owner_lookup_key__in=normalized_names)
        .values_list("owner_lookup_key", flat=True)
        .distinct()
    )


def _attach_certificate_owner_metadata(rows):
    owner_names = [row.get("owner_name") for row in rows if row.get("owner_name")]
    owner_profile_lookup = _build_owner_profile_lookup(owner_names)
    manual_owner_preview_keys = _build_certificate_manual_owner_preview_keys(owner_names)
    default_owner_profile_image_url = static("images/default-user-image.jpg")

    for row in rows:
        owner_name = " ".join((row.get("owner_name") or "").split()).strip() or "Unknown Owner"
        normalized_owner = _normalize_person_name(owner_name)
        matched_owner_profile = owner_profile_lookup.get(normalized_owner, {})
        matched_owner_user_id = matched_owner_profile.get("user_id")

        row["owner_name"] = owner_name
        row["profile_image_url"] = (
            matched_owner_profile.get("image_url")
            or default_owner_profile_image_url
        )
        row["owner_initials"] = _owner_initials(owner_name)
        row["owner_type_label"] = "User" if matched_owner_user_id else "Owner"

        if matched_owner_user_id:
            row["owner_group_key"] = f"user:{matched_owner_user_id}"
            row["profile_url"] = reverse(
                "dogadoption_admin:registration_owner_profile",
                args=[matched_owner_user_id],
            )
        else:
            owner_key = normalized_owner or f"registration:{row.get('registration_id', 'unknown')}"
            row["owner_group_key"] = f"name:{owner_key}"
            if normalized_owner in manual_owner_preview_keys:
                row["profile_url"] = (
                    f"{reverse('dogadoption_admin:registration_owner_profile', args=[0])}"
                    f"?{urlencode({'owner_key': normalized_owner, 'owner_name': owner_name})}"
                )
            else:
                row["profile_url"] = ""

        row["profile_available"] = bool(row["profile_url"])

    return rows


def _build_certificate_owner_groups(rows):
    grouped_rows = defaultdict(list)
    for row in rows:
        grouped_rows[row["owner_group_key"]].append(row)

    owner_groups = []
    for owner_group_key, items in grouped_rows.items():
        sorted_items = sorted(
            items,
            key=lambda item: (
                _certificate_reg_no_sort_key(item.get("reg_no")),
                item.get("vaccination_date") or datetime.min.date(),
                item.get("registration_id") or 0,
            ),
            reverse=True,
        )
        first_item = sorted_items[0]
        unique_registration_ids = {
            item.get("registration_id")
            for item in sorted_items
            if item.get("registration_id") is not None
        }
        owner_groups.append(
            {
                "owner_group_key": owner_group_key,
                "owner_name": first_item.get("owner_name") or "Unknown Owner",
                "profile_image_url": first_item.get("profile_image_url") or "",
                "owner_initials": first_item.get("owner_initials") or "?",
                "profile_url": first_item.get("profile_url") or "",
                "profile_available": bool(first_item.get("profile_url")),
                "owner_type_label": first_item.get("owner_type_label") or "Owner",
                "vaccinated_pet_count": len(unique_registration_ids),
                "vaccination_card_count": len(sorted_items),
                "rows": sorted_items,
            }
        )

    owner_groups.sort(
        key=lambda group: (
            -group["vaccinated_pet_count"],
            -group["vaccination_card_count"],
            group["owner_name"].casefold(),
            group["owner_group_key"],
        )
    )
    return owner_groups


def _format_cert_date(value):
    if not value:
        return ""
    if hasattr(value, "strftime"):
        return value.strftime("%m-%d-%Y")
    return str(value)


def _pad_rows(rows, min_rows):
    padded = list(rows)
    while len(padded) < min_rows:
        padded.append({})
    return padded


def _build_certificate_payload(
    registration,
    vaccinations=None,
    dewormings=None,
    vac_limit=10,
    vac_min_rows=10,
    dew_limit=8,
    dew_min_rows=8,
):
    vac_records = list(vaccinations) if vaccinations is not None else list(
        VaccinationRecord.objects.filter(registration=registration).order_by("-date")
    )
    dew_records = list(dewormings) if dewormings is not None else list(
        DewormingTreatmentRecord.objects.filter(registration=registration).order_by("-date")
    )

    vac_rows = []
    for record in vac_records[:vac_limit]:
        vac_rows.append({
            "date": _format_cert_date(record.date),
            "vaccine_name": record.vaccine_name or "",
            "manufacturer_lot_no": record.manufacturer_lot_no or "",
            "vaccine_expiry_date": _format_cert_date(record.vaccine_expiry_date),
            "vaccination_expiry_date": _format_cert_date(record.vaccination_expiry_date),
            "veterinarian": record.veterinarian or "",
        })

    dew_rows = []
    for record in dew_records[:dew_limit]:
        route = (record.route or "").strip()
        frequency = (record.frequency or "").strip()
        route_frequency = f"{route} / {frequency}".strip(" /") if (route or frequency) else ""
        dew_rows.append({
            "date": _format_cert_date(record.date),
            "medicine_given": record.medicine_given or "",
            "route_frequency": route_frequency,
            "veterinarian": record.veterinarian or "",
        })

    return {
        "id": registration.id,
        "reg_no": registration.reg_no or "",
        "name_of_pet": registration.name_of_pet or "",
        "breed": registration.breed or "",
        "dob": _format_cert_date(registration.dob),
        "color_markings": registration.color_markings or "",
        "sex": registration.sex or "",
        "is_male": registration.sex == "M",
        "is_female": registration.sex == "F",
        "status": registration.status or "",
        "is_castrated": registration.status == "Castrated",
        "is_spayed": registration.status == "Spayed",
        "is_intact": registration.status == "Intact",
        "owner_name": registration.owner_name or "",
        "address": registration.address or "",
        "contact_no": registration.contact_no or "",
        "vaccination_rows": _pad_rows(vac_rows, vac_min_rows),
        "deworming_rows": _pad_rows(dew_rows, dew_min_rows),
        "vaccination_count": len(vac_records),
        "deworming_count": len(dew_records),
        "has_vaccinations": bool(vac_records),
        "has_dewormings": bool(dew_records),
    }


def _get_profile_or_none(user):
    if user is None:
        return None
    try:
        return user.profile
    except Profile.DoesNotExist:
        return None


def _dog_capture_request_board_queryset():
    return (
        DogCaptureRequest.objects.select_related(
            "requested_by",
            "requested_by__profile",
            "assigned_admin",
        )
        .only(
            "id",
            "requested_by_id",
            "assigned_admin_id",
            "request_type",
            "submission_type",
            "preferred_appointment_date",
            "reason",
            "description",
            "latitude",
            "longitude",
            "barangay",
            "city",
            "manual_full_address",
            "image",
            "status",
            "scheduled_date",
            "captured_at",
            "admin_message",
            "created_at",
            "requested_by__id",
            "requested_by__username",
            "requested_by__first_name",
            "requested_by__last_name",
            "requested_by__profile__address",
            "requested_by__profile__phone_number",
            "requested_by__profile__facebook_url",
            "requested_by__profile__profile_image",
            "assigned_admin__id",
            "assigned_admin__username",
            "assigned_admin__first_name",
            "assigned_admin__last_name",
        )
        .order_by("-created_at", "-id")
    )


def _dog_capture_request_map_queryset():
    return (
        DogCaptureRequest.objects.select_related(
            "requested_by",
            "requested_by__profile",
        )
        .only(
            "id",
            "requested_by_id",
            "request_type",
            "submission_type",
            "reason",
            "latitude",
            "longitude",
            "barangay",
            "city",
            "manual_full_address",
            "image",
            "status",
            "created_at",
            "requested_by__username",
            "requested_by__first_name",
            "requested_by__last_name",
            "requested_by__profile__address",
            "requested_by__profile__phone_number",
        )
        .order_by("-created_at", "-id")
    )


def _dog_capture_request_first_image_urls(request_ids):
    if not request_ids:
        return {}

    image_urls = {}
    for request_image in (
        DogCaptureRequestImage.objects.filter(request_id__in=request_ids)
        .only("request_id", "image")
        .order_by("id")
    ):
        if request_image.request_id in image_urls:
            continue
        image_url = _safe_media_url(request_image.image)
        if image_url:
            image_urls[request_image.request_id] = image_url
    return image_urls


def _enrich_capture_request_user(req):
    user = req.requested_by
    profile = _get_profile_or_none(user)

    name_parts = [user.first_name, user.last_name]
    full_name = " ".join(part for part in name_parts if part).strip() or user.username

    req.requester_full_name = full_name
    req.requester_phone = (
        profile.phone_number.strip()
        if profile and profile.phone_number
        else "No phone number"
    )
    req.requester_address = (
        profile.address.strip()
        if profile and profile.address
        else "No address provided"
    )
    req.requester_facebook = (
        profile.facebook_url.strip()
        if profile and profile.facebook_url
        else ""
    )


def _enrich_capture_request_display(req):
    _enrich_capture_request_user(req)

    barangay = _clean_barangay(req.barangay)
    city = _clean_barangay(req.city)
    manual_full_address = _clean_barangay(req.manual_full_address)
    profile = _get_profile_or_none(req.requested_by)
    profile_address = _clean_barangay(getattr(profile, "address", ""))

    if not barangay:
        profile_barangay = profile_address
        barangay = _resolve_barangay_name(profile_barangay) or profile_barangay

    if not city:
        city = _extract_city_from_address(profile_address)

    if manual_full_address:
        req.location_label = manual_full_address
    elif barangay and city:
        req.location_label = f"{barangay}, {city}"
    elif barangay:
        req.location_label = barangay
    elif city:
        req.location_label = city
    elif req.latitude is not None and req.longitude is not None:
        req.location_label = "Pinned location"
    else:
        req.location_label = "No location"
    req.has_location = req.location_label != "No location"
    req.display_barangay = barangay


ANNOUNCEMENT_CATEGORY_OPTIONS = [
    {
        "slug": "dog-announcements",
        "value": DogAnnouncement.CATEGORY_DOG_ANNOUNCEMENT,
        "label": "Dog Announcements",
        "description": (
            "For vaccination programs, educational campaigns, dog-related events, "
            "and general dog care information."
        ),
        "topics": [
            "Vaccination programs",
            "Educational campaigns",
            "Dog-related events",
            "General dog care information",
        ],
    },
    {
        "slug": "dog-laws",
        "value": DogAnnouncement.CATEGORY_DOG_LAW,
        "label": "Dog Laws",
        "description": (
            "For rules and regulations about dogs, local ordinances, and legal "
            "responsibilities of dog owners."
        ),
        "topics": [
            "Rules and regulations about dogs",
            "Local ordinances",
            "Legal responsibilities of dog owners",
        ],
    },
]

ANNOUNCEMENT_CATEGORY_BY_SLUG = {
    option["slug"]: option for option in ANNOUNCEMENT_CATEGORY_OPTIONS
}

ANNOUNCEMENT_CATEGORY_BY_VALUE = {
    option["value"]: option for option in ANNOUNCEMENT_CATEGORY_OPTIONS
}

ANNOUNCEMENT_BUCKET_VALUES = {
    value for value, _label in DogAnnouncement.DISPLAY_BUCKET_CHOICES
}

ADMIN_ANNOUNCEMENT_PAGE_SIZE = 24
AUTOCOMPLETE_RESULTS_DEFAULT = 12
AUTOCOMPLETE_RESULTS_MAX = 25
ANALYTICS_DASHBOARD_CACHE_TTL_SECONDS = 60


def _parse_positive_int(raw_value, default=AUTOCOMPLETE_RESULTS_DEFAULT, max_value=AUTOCOMPLETE_RESULTS_MAX):
    try:
        parsed = int(raw_value)
    except (TypeError, ValueError):
        parsed = default
    return max(1, min(parsed, max_value))


def _cacheable_json_response(payload, *, max_age, public=False, status=200):
    response = JsonResponse(payload, status=status)
    visibility = "public" if public else "private"
    response["Cache-Control"] = f"{visibility}, max-age={max_age}"
    return response


def _set_post_form_barangay_source(post_form):
    return post_form


def _is_truthy(value):
    return str(value).strip().lower() in {"1", "true", "yes", "on"}


def _safe_next_url(request):
    next_url = (request.POST.get("next") or request.GET.get("next") or "").strip()
    if next_url and url_has_allowed_host_and_scheme(
        next_url,
        allowed_hosts={request.get_host()},
        require_https=request.is_secure(),
    ):
        return next_url
    return ""


def _redirect_to_safe_next(request, fallback_name, *args):
    next_url = _safe_next_url(request)
    if next_url:
        return redirect(next_url)
    return redirect(fallback_name, *args)


def _touch_post_board_cache():
    cache.delete(POST_HISTORY_CACHE_KEY)
    bump_user_home_feed_namespace()
    invalidate_user_notification_content()


def _annotate_post_request_count(qs, annotation_name, request_type=None):
    request_qs = PostRequest.objects.filter(post_id=OuterRef("pk"))
    if request_type:
        request_qs = request_qs.filter(request_type=request_type)
    request_count_subquery = (
        request_qs.order_by()
        .values("post_id")
        .annotate(total=Count("id"))
        .values("total")[:1]
    )
    return qs.annotate(
        **{
            annotation_name: Coalesce(
                Subquery(request_count_subquery, output_field=IntegerField()),
                Value(0),
            )
        }
    )


def _build_post_form_page_context(post_form, *, post=None, form_action="", form_mode="create", back_url=""):
    _set_post_form_barangay_source(post_form)
    existing_images = []
    if post is not None:
        for image in post.images.only("image").order_by("id"):
            image_url = _safe_media_url(getattr(image, "image", None))
            if image_url:
                existing_images.append(image_url)
    is_edit_mode = form_mode == "edit"
    return {
        "post_form": post_form,
        "post": post,
        "form_action": form_action,
        "form_mode": form_mode,
        "form_title": "Edit Dog Post" if is_edit_mode else "Create Dog Adoption Post",
        "submit_label": "Save Changes" if is_edit_mode else "Publish Post",
        "submit_icon": "fas fa-save" if is_edit_mode else "fas fa-paper-plane",
        "confirm_title": "Confirm Save" if is_edit_mode else "Confirm Publish",
        "confirm_body": (
            "Save these changes to the post now?"
            if is_edit_mode
            else "Publish this post now?"
        ),
        "confirm_submit_label": "Save" if is_edit_mode else "Publish",
        "back_url": back_url or reverse("dogadoption_admin:post_list"),
        "require_images": not is_edit_mode,
        "existing_images": existing_images,
    }


def _parse_appointment_dates(dates_raw):
    parsed_dates = []
    for value in [v.strip() for v in (dates_raw or "").split(",") if v.strip()]:
        parsed_date = parse_date(value)
        if parsed_date:
            parsed_dates.append(parsed_date)
    return sorted(set(parsed_dates))


def _save_global_appointment_dates(parsed_dates, user):
    today = timezone.localdate()
    editable_dates = [day for day in parsed_dates if day >= today]
    with transaction.atomic():
        GlobalAppointmentDate.objects.filter(
            appointment_date__gte=today,
        ).exclude(
            appointment_date__in=editable_dates
        ).delete()
        for day in editable_dates:
            GlobalAppointmentDate.objects.update_or_create(
                appointment_date=day,
                defaults={
                    "created_by": user,
                    "is_active": True,
                },
            )


def _validate_and_save_global_appointment_dates(dates_raw, user):
    parsed_dates = _parse_appointment_dates(dates_raw)
    today = timezone.localdate()
    submitted_past_dates = {day for day in parsed_dates if day < today}
    if submitted_past_dates:
        locked_past_dates = set(
            GlobalAppointmentDate.objects.filter(
                appointment_date__lt=today
            ).values_list("appointment_date", flat=True)
        )
        if not submitted_past_dates.issubset(locked_past_dates):
            return False
    _save_global_appointment_dates(parsed_dates, user)
    return True


def _get_active_global_appointment_dates():
    return list(
        GlobalAppointmentDate.objects.filter(
            is_active=True
        ).order_by("appointment_date").values_list("appointment_date", flat=True)
    )


def _get_available_appointment_dates():
    return GlobalAppointmentDate.objects.filter(
        is_active=True,
        appointment_date__gte=timezone.localdate(),
    ).order_by("appointment_date")


def _build_requests_with_meta(post, request_type):
    requests_qs = (
        post.requests.filter(request_type=request_type)
        .select_related("user", "post")
        .prefetch_related("images")
        .order_by("-created_at")
    )
    requests = list(requests_qs)
    user_ids = [req.user_id for req in requests]

    profiles = Profile.objects.filter(user_id__in=user_ids)
    profile_by_user_id = {profile.user_id: profile for profile in profiles}

    requests_with_meta = []
    for req in requests:
        requests_with_meta.append({
            "req": req,
            "profile": profile_by_user_id.get(req.user_id),
        })
    return requests_with_meta


def _build_request_redirect(req):
    if req.request_type == "claim":
        return redirect("dogadoption_admin:claim_requests", req.post.id)
    return redirect("dogadoption_admin:adoption_requests", req.post.id)


def _build_request_redirect_or_next(request, req):
    next_url = (request.POST.get("next") or "").strip()
    if next_url and url_has_allowed_host_and_scheme(
        next_url,
        allowed_hosts={request.get_host()},
        require_https=request.is_secure(),
    ):
        return redirect(next_url)
    return _build_request_redirect(req)


def _render_post_request_list(request, post_id, request_type, template_name):
    post = get_object_or_404(Post.objects.filter(is_history=False), id=post_id)
    return render(request, template_name, {
        "post": post,
        "requests_meta": _build_requests_with_meta(post, request_type),
        "available_dates": _get_available_appointment_dates(),
    })


def _apply_registration_date_filter(dogs, date_filter_type, filter_date, filter_month, filter_year):
    date_filter_label = ""

    if date_filter_type == "day" and filter_date:
        selected_date = parse_date(filter_date)
        if selected_date:
            dogs = dogs.filter(date_registered=selected_date)
            date_filter_label = selected_date.strftime("%b %d, %Y")
        else:
            date_filter_type = "all"
    elif date_filter_type == "month" and filter_month:
        month_match = re.match(r"^(\d{4})-(\d{2})$", filter_month)
        if month_match:
            month_year = int(month_match.group(1))
            month_value = int(month_match.group(2))
            if 1 <= month_value <= 12:
                dogs = dogs.filter(
                    date_registered__year=month_year,
                    date_registered__month=month_value,
                )
                date_filter_label = datetime(month_year, month_value, 1).strftime("%B %Y")
            else:
                date_filter_type = "all"
        else:
            date_filter_type = "all"
    elif date_filter_type == "year" and filter_year:
        if filter_year.isdigit():
            year_value = int(filter_year)
            if 1900 <= year_value <= 9999:
                dogs = dogs.filter(date_registered__year=year_value)
                date_filter_label = str(year_value)
            else:
                date_filter_type = "all"
        else:
            date_filter_type = "all"
    else:
        date_filter_type = "all"

    return dogs, date_filter_type, date_filter_label


def _build_registration_filter_params(date_filter_type, filter_date, filter_month, filter_year):
    if date_filter_type == "day" and filter_date:
        return {
            "date_filter_type": "day",
            "filter_date": filter_date,
        }
    if date_filter_type == "month" and filter_month:
        return {
            "date_filter_type": "month",
            "filter_month": filter_month,
        }
    if date_filter_type == "year" and filter_year:
        return {
            "date_filter_type": "year",
            "filter_year": filter_year,
        }
    return {}


def _create_vaccination_and_update_defaults(
    registration,
    cert_settings,
    vac_date,
    vaccine_name,
    manufacturer_lot_no,
    vaccine_expiry_date,
    vaccination_expiry_date,
):
    VaccinationRecord.objects.create(
        registration=registration,
        date=vac_date,
        vaccine_name=vaccine_name,
        manufacturer_lot_no=manufacturer_lot_no,
        vaccine_expiry_date=vaccine_expiry_date,
        vaccination_expiry_date=vaccination_expiry_date,
        veterinarian="",
    )
    settings_obj = cert_settings or CertificateSettings.objects.create()
    if vac_date:
        parsed_vac_date = parse_date(vac_date)
        if parsed_vac_date:
            settings_obj.default_vac_date = parsed_vac_date
    if vaccine_name:
        settings_obj.default_vaccine_name = vaccine_name
    if manufacturer_lot_no:
        settings_obj.default_manufacturer_lot_no = manufacturer_lot_no
    if vaccine_expiry_date:
        parsed_vac_expiry = parse_date(vaccine_expiry_date)
        if parsed_vac_expiry:
            settings_obj.default_vaccine_expiry_date = parsed_vac_expiry
    settings_obj.save(update_fields=[
        "default_vac_date",
        "default_vaccine_name",
        "default_manufacturer_lot_no",
        "default_vaccine_expiry_date",
    ])
    return settings_obj


def _get_vaccination_post_values(request):
    return (
        (request.POST.get("vac_date") or "").strip(),
        (request.POST.get("vaccine_name") or "").strip(),
        (request.POST.get("manufacturer_lot_no") or "").strip(),
        (request.POST.get("vaccine_expiry_date") or "").strip(),
        (request.POST.get("vaccination_expiry_date") or "").strip(),
    )


def _latest_certificate_record_date(rows):
    return next((row.get("date", "") for row in rows if row.get("date")), "")


DOG_REGISTRATION_MAX_IMAGES = 12
DOG_REGISTRATION_MAX_IMAGE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB
DOG_REGISTRATION_OWNER_MAX_PETS = 4


def _validate_registration_images(uploaded_images):
    if len(uploaded_images) > DOG_REGISTRATION_MAX_IMAGES:
        return f"You can upload up to {DOG_REGISTRATION_MAX_IMAGES} photos per registration."

    for image in uploaded_images:
        if image.size <= 0:
            return "One of the uploaded files is empty."
        if image.size > DOG_REGISTRATION_MAX_IMAGE_SIZE_BYTES:
            return "Each image must be 10 MB or smaller."

        content_type = (getattr(image, "content_type", "") or "").lower()
        if not content_type.startswith("image/"):
            return "Only image files are allowed for registration photos."

    return ""


# =============================================================================
# Shared admin access and authentication helpers
# =============================================================================

def _is_ajax_request(request):
    """Return True when the request was sent through frontend fetch/XHR code."""
    return request.headers.get("x-requested-with") == "XMLHttpRequest"

def admin_required(view_func):
    """Limit a view to authenticated staff users."""

    @wraps(view_func)
    def wrapper(request, *args, **kwargs):
        if not request.user.is_authenticated:
            if _is_ajax_request(request):
                return JsonResponse({
                    "ok": False,
                    "auth_required": True,
                    "auth_modal": "login",
                    "login_url": reverse("user:login"),
                }, status=401)
            return redirect('user:login')
        if not request.user.is_staff:
            redirect_url = reverse("user:user_home")
            if _is_ajax_request(request):
                return JsonResponse({
                    "ok": False,
                    "redirect_url": redirect_url,
                }, status=403)
            return redirect(redirect_url)
        access = get_admin_access(request.user)
        route_name = getattr(getattr(request, "resolver_match", None), "url_name", "") or view_func.__name__
        if not is_route_allowed(access, route_name):
            redirect_url = access.get("landing_url") or reverse("dogadoption_admin:admin_edit_profile")
            if _is_ajax_request(request):
                return JsonResponse({
                    "ok": False,
                    "message": "You do not have access to that section.",
                    "redirect_url": redirect_url,
                }, status=403)
            messages.error(request, "You do not have access to that section.")
            return redirect(redirect_url)
        request.admin_access = access
        return view_func(request, *args, **kwargs)

    return wrapper


def admin_login(request):
    """Reuse the main user login page for admin authentication."""
    return redirect('user:login')


@login_required
@require_POST
def admin_logout(request):
    """Log the admin out and clear the dedicated admin session cookie."""
    logout(request)
    response = redirect(
        f"{reverse('user:user_home')}?feed_token={secrets.token_hex(8)}"
    )
    response.delete_cookie('admin_sessionid')
    return response


# =============================================================================
# Navigation 1/5: Home
# Covers the Home sidebar link, including posts and adoption/claim review.
# =============================================================================

@admin_required
def create_post(request):
    """Create a rescue post from the admin home screen."""
    access = getattr(request, "admin_access", get_admin_access(request.user))
    if not access.get("can_create_posts"):
        messages.error(request, "You do not have permission to create posts.")
        return redirect(access.get("landing_url") or reverse("dogadoption_admin:admin_edit_profile"))
    if request.method == 'POST':
        post_form = PostForm(request.POST)

        if post_form.is_valid():
            post = post_form.save(commit=False)
            post.user = request.user
            post.save()

            # Multiple images
            for image in request.FILES.getlist('images'):
                PostImage.objects.create(post=post, image=image)

            _touch_post_board_cache()
            messages.success(request, "Post created successfully.", extra_tags="post_list")
            return redirect('dogadoption_admin:post_list')
    else:
        post_form = PostForm()

    return render(
        request,
        'admin_home/create_post.html',
        _build_post_form_page_context(
            post_form,
            form_action=reverse("dogadoption_admin:create_post"),
            form_mode="create",
            back_url=reverse("dogadoption_admin:post_list"),
        ),
    )


@admin_required
@require_http_methods(["GET", "POST"])
def update_post(request, post_id):
    """Edit an existing rescue post from the admin home module."""
    access = getattr(request, "admin_access", get_admin_access(request.user))
    if not access.get("can_create_posts"):
        messages.error(request, "You do not have permission to edit posts.")
        return redirect(access.get("landing_url") or reverse("dogadoption_admin:post_list"))

    post = get_object_or_404(Post.objects.prefetch_related("images").filter(is_history=False), id=post_id)

    if request.method == "POST":
        post_form = PostForm(request.POST, instance=post)
        if post_form.is_valid():
            post = post_form.save()
            for image in request.FILES.getlist("images"):
                PostImage.objects.create(post=post, image=image)
            _touch_post_board_cache()
            messages.success(request, "Post updated successfully.", extra_tags="post_list")
            return _redirect_to_safe_next(request, "dogadoption_admin:post_list")
    else:
        post_form = PostForm(instance=post)

    return render(
        request,
        "admin_home/create_post.html",
        _build_post_form_page_context(
            post_form,
            post=post,
            form_action=reverse("dogadoption_admin:update_post", args=[post.id]),
            form_mode="edit",
            back_url=_safe_next_url(request) or reverse("dogadoption_admin:post_list"),
        ),
    )


@admin_required
@require_POST
def toggle_post_pin(request, post_id):
    """Pin or unpin a rescue post for the public home spotlight."""
    access = getattr(request, "admin_access", get_admin_access(request.user))
    if not access.get("can_create_posts"):
        messages.error(request, "You do not have permission to pin posts.", extra_tags="post_list")
        return _redirect_to_safe_next(request, "dogadoption_admin:post_list")

    post = get_object_or_404(Post.objects.filter(is_history=False), id=post_id)
    if post.status not in {"rescued", "under_care"}:
        messages.warning(
            request,
            "Only active rescue posts can be pinned to the homepage.",
            extra_tags="post_list",
        )
        return _redirect_to_safe_next(request, "dogadoption_admin:post_list")

    post_name = post.display_title or post.caption or f"Post {post.id}"
    if post.is_pinned:
        post.is_pinned = False
        post.pinned_at = None
        post.save(update_fields=["is_pinned", "pinned_at"])
        messages.success(
            request,
            f'Post "{post_name}" was removed from the homepage spotlight.',
            extra_tags="post_list",
        )
    else:
        now = timezone.now()
        pinned_count = Post.objects.filter(is_pinned=True).exclude(pk=post.pk).count()
        if pinned_count >= HOME_SPOTLIGHT_PIN_LIMIT:
            messages.warning(
                request,
                f"Only {HOME_SPOTLIGHT_PIN_LIMIT} dogs can be pinned at once.",
                extra_tags="post_list",
            )
            return _redirect_to_safe_next(request, "dogadoption_admin:post_list")
        with transaction.atomic():
            post.is_pinned = True
            post.pinned_at = now
            post.save(update_fields=["is_pinned", "pinned_at"])
        messages.success(
            request,
            f'Post "{post_name}" is now pinned on the homepage spotlight.',
            extra_tags="post_list",
        )

    _touch_post_board_cache()
    return _redirect_to_safe_next(request, "dogadoption_admin:post_list")


@admin_required
@require_POST
def toggle_post_phase(request, post_id):
    """Manually move an active rescue post between the redeem and adoption groups."""
    access = getattr(request, "admin_access", get_admin_access(request.user))
    if not access.get("can_create_posts"):
        messages.error(request, "You do not have permission to update posts.", extra_tags="post_list")
        return _redirect_to_safe_next(request, "dogadoption_admin:post_list")

    post = get_object_or_404(Post.objects.filter(is_history=False), id=post_id)
    if post.status in {"reunited", "adopted"}:
        messages.warning(request, "Finalized posts can no longer switch between redeem and adoption.", extra_tags="post_list")
        return _redirect_to_safe_next(request, "dogadoption_admin:post_list")

    timeline_phase = post.timeline_phase()
    if timeline_phase == "closed":
        messages.warning(request, "This post is already outside the active redeem/adoption window.", extra_tags="post_list")
        return _redirect_to_safe_next(request, "dogadoption_admin:post_list")

    current_phase = post.current_phase()
    target_phase = (request.POST.get("phase") or "").strip().lower()
    if target_phase not in {"claim", "adopt"}:
        if current_phase == "claim":
            target_phase = "adopt"
        elif current_phase == "adopt":
            target_phase = "claim"

    if target_phase not in {"claim", "adopt"}:
        messages.error(request, "Invalid target status selected.", extra_tags="post_list")
        return _redirect_to_safe_next(request, "dogadoption_admin:post_list")

    post.phase_override = target_phase
    post.phase_override_started_at = timezone.now()
    post.save(update_fields=["phase_override", "phase_override_started_at"])
    _touch_post_board_cache()
    messages.success(
        request,
        f'Post moved to {"For Redeem" if target_phase == "claim" else "For Adoption"}.',
        extra_tags="post_list",
    )
    return _redirect_to_safe_next(request, "dogadoption_admin:post_list")


@admin_required
@require_POST
def finalize_post(request, post_id):
    """Manually complete an active rescue post and remove it from the live board."""
    access = getattr(request, "admin_access", get_admin_access(request.user))
    if not access.get("can_create_posts"):
        messages.error(request, "You do not have permission to finalize posts.", extra_tags="post_list")
        return _redirect_to_safe_next(request, "dogadoption_admin:post_list")

    with transaction.atomic():
        post = get_object_or_404(
            Post.objects.select_for_update().filter(is_history=False),
            id=post_id,
        )
        post_name = post.display_title or post.caption or f"Post {post.id}"

        if post.status in {"reunited", "adopted"}:
            messages.warning(request, "This post has already been finalized.", extra_tags="post_list")
            return _redirect_to_safe_next(request, "dogadoption_admin:post_list")

        current_phase = post.current_phase()
        if current_phase not in {"claim", "adopt"}:
            messages.warning(
                request,
                "Only active redeem or adoption posts can be finalized from the board.",
                extra_tags="post_list",
            )
            return _redirect_to_safe_next(request, "dogadoption_admin:post_list")

        expected_status = "reunited" if current_phase == "claim" else "adopted"
        target_status = (request.POST.get("status") or "").strip().lower() or expected_status
        if target_status != expected_status:
            messages.error(request, "Invalid finalize action selected for this post.", extra_tags="post_list")
            return _redirect_to_safe_next(request, "dogadoption_admin:post_list")

        post.status = target_status
        post.phase_override = ""
        post.phase_override_started_at = None
        post.is_pinned = False
        post.pinned_at = None
        post.save(
            update_fields=[
                "status",
                "phase_override",
                "phase_override_started_at",
                "is_pinned",
                "pinned_at",
            ]
        )

        pending_requests = list(post.requests.filter(status="pending").only("id", "user_id"))
        if pending_requests:
            post.requests.filter(id__in=[req.id for req in pending_requests]).update(
                status="rejected",
                scheduled_appointment_date=None,
            )
            for pending_req in pending_requests:
                invalidate_user_notification_payload(pending_req.user_id)

    _touch_post_board_cache()
    messages.success(
        request,
        (
            f'Post "{post_name}" was marked as redeemed and moved out of the active board.'
            if target_status == "reunited"
            else f'Post "{post_name}" was marked as adopted and moved out of the active board.'
        ),
        extra_tags="post_list",
    )
    return _redirect_to_safe_next(request, "dogadoption_admin:post_list")


@admin_required
@require_POST
def delete_post(request, post_id):
    """Archive a rescue post or permanently remove it from the system."""
    access = getattr(request, "admin_access", get_admin_access(request.user))
    if not access.get("can_create_posts"):
        messages.error(request, "You do not have permission to delete posts.", extra_tags="post_list")
        return _redirect_to_safe_next(request, "dogadoption_admin:post_list")

    post = get_object_or_404(Post, id=post_id)
    post_name = post.display_title or post.caption or f"Post {post.id}"
    permanent_delete = _is_truthy(request.POST.get("permanent_delete"))

    if permanent_delete:
        post.delete()
        messages.success(request, f'Post "{post_name}" was permanently deleted.', extra_tags="post_list")
    else:
        if not post.is_history:
            post.is_history = True
            post.is_pinned = False
            post.pinned_at = None
            post.save(update_fields=["is_history", "is_pinned", "pinned_at"])
        messages.success(request, f'Post "{post_name}" was moved to history.', extra_tags="post_list")

    _touch_post_board_cache()
    return _redirect_to_safe_next(request, "dogadoption_admin:post_list")

@admin_required
def post_list(request):
    """Render the post board and handle quick-create or appointment updates."""
    access = getattr(request, "admin_access", get_admin_access(request.user))
    show_create_modal = False
    show_appointment_modal = request.method == "GET" and (
        request.GET.get("open_appointment", "").lower() in {"1", "true", "yes"}
    )
    if show_appointment_modal and not access.get("is_full_admin"):
        messages.error(request, "Only the admin can set appointment dates.")
        show_appointment_modal = False
    post_form = PostForm()
    if request.method == 'POST':
        form_type = (request.POST.get("form_type") or "").strip()

        if form_type == "appointment_dates":
            if not access.get("is_full_admin"):
                messages.error(request, "Only the admin can set appointment dates.", extra_tags="post_list")
                return redirect(reverse("dogadoption_admin:post_list"))
            show_appointment_modal = True
            dates_raw = (request.POST.get('appointment_dates') or '').strip()
            if not _validate_and_save_global_appointment_dates(dates_raw, request.user):
                messages.error(
                    request,
                    "Past appointment dates are locked and cannot be changed.",
                    extra_tags="post_list",
                )
            else:
                messages.success(request, "Appointment dates saved.", extra_tags="post_list")
                return redirect(reverse('dogadoption_admin:post_list'))
        else:
            if not access.get("can_create_posts"):
                messages.error(request, "You do not have permission to create posts.", extra_tags="post_list")
                return redirect(reverse("dogadoption_admin:post_list"))
            post_form = PostForm(request.POST)
            show_create_modal = True

            if post_form.is_valid():
                post = post_form.save(commit=False)
                post.user = request.user
                post.save()

                for image in request.FILES.getlist('images'):
                    PostImage.objects.create(post=post, image=image)

                _touch_post_board_cache()
                messages.success(request, "Post created successfully.", extra_tags="post_list")
                return redirect(reverse('dogadoption_admin:post_list'))

    _set_post_form_barangay_source(post_form)

    now = timezone.now()
    active_statuses = ["rescued", "under_care"]
    rows_per_page = 8
    section_definitions = (
        {
            "key": "claim",
            "title": "For Redeem",
            "tone": "status-claim",
            "page_param": "claim_page",
            "request_type": "claim",
            "allow_toggle": True,
            "empty_message": "No posts are currently queued for redeem.",
        },
        {
            "key": "adopt",
            "title": "For Adoption",
            "tone": "status-adopt",
            "page_param": "adoption_page",
            "request_type": "adopt",
            "allow_toggle": True,
            "empty_message": "No posts are currently queued for adoption.",
        },
    )

    base_qs = Post.with_pending_request_state(
        Post.objects.filter(is_history=False).only(
            'id',
            'caption',
            'breed',
            'breed_other',
            'age_group',
            'size_group',
            'gender',
            'coat_length',
            'colors',
            'color_other',
            'location',
            'status',
            'rescued_date',
            'claim_days',
            'created_at',
            'phase_override',
            'phase_override_started_at',
            'is_history',
            'is_pinned',
            'pinned_at',
            'view_count',
        )
    )
    base_qs = _annotate_post_request_count(base_qs, "claim_count", "claim")
    base_qs = _annotate_post_request_count(base_qs, "adopt_count", "adopt")
    base_qs = _annotate_post_request_count(base_qs, "total_request_count")

    def _build_page_qs(page_param, page_num):
        params = request.GET.copy()
        params[page_param] = str(page_num)
        return params.urlencode()

    def _build_post_item(post, section_key, request_type, allow_toggle):
        days = hours = minutes = 0
        is_active_section = section_key in {"claim", "adopt"}
        is_pending_review = (
            is_active_section
            and bool(getattr(post, f"has_pending_{section_key}_request", False))
        )
        pending_review_until = (
            post.pending_request_review_available_at(section_key)
            if is_pending_review
            else None
        )
        pending_review_until_label = (
            timezone.localtime(pending_review_until).strftime("%b %d, %Y %I:%M %p")
            if pending_review_until
            else ""
        )
        deadline = None
        if section_key == "claim" and not is_pending_review:
            deadline = post.claim_deadline()
        elif section_key == "adopt" and not is_pending_review:
            deadline = post.adoption_deadline()

        remaining_time = timedelta(seconds=0)
        if is_active_section and not is_pending_review:
            if deadline:
                remaining_time = deadline - now
            total_seconds = max(int(remaining_time.total_seconds()), 0)
            days = total_seconds // 86400
            remainder = total_seconds % 86400
            hours = remainder // 3600
            remainder = remainder % 3600
            minutes = remainder // 60

        time_left_label = "Completed"
        if is_pending_review:
            time_left_label = (
                f"Verification until {pending_review_until_label}"
                if pending_review_until_label
                else "Pending admin review"
            )
        elif is_active_section:
            time_left_label = (
                f"{days:02d}d {hours:02d}h {minutes:02d}m"
                if deadline
                else "No active time window"
            )

        toggle_phase = ""
        toggle_label = ""
        if allow_toggle:
            toggle_phase = "adopt" if section_key == "claim" else "claim"
            toggle_label = "Switch to Adoption" if section_key == "claim" else "Switch to Redeem"

        request_count = int(getattr(post, "total_request_count", 0) or 0)
        if request_type == "claim":
            request_count = int(getattr(post, "claim_count", 0) or 0)
        elif request_type == "adopt":
            request_count = int(getattr(post, "adopt_count", 0) or 0)

        return {
            "post": post,
            "section_key": section_key,
            "request_type": request_type,
            "request_count": request_count,
            "total_request_count": int(getattr(post, "total_request_count", 0) or 0),
            "claim_request_count": int(getattr(post, "claim_count", 0) or 0),
            "adopt_request_count": int(getattr(post, "adopt_count", 0) or 0),
            "view_count": int(getattr(post, "view_count", 0) or 0),
            "days_left": days,
            "hours_left": hours,
            "minutes_left": minutes,
            "is_pending_review": is_pending_review,
            "pending_review_until": pending_review_until,
            "pending_review_until_label": pending_review_until_label,
            "deadline_iso": deadline.isoformat() if deadline else "",
            "time_left_label": time_left_label,
            "is_pinned": bool(getattr(post, "is_pinned", False)),
            "claim_requests": [],
            "adopt_requests": [],
            "primary_image_url": "",
            "finalized_user_name": "-",
            "finalized_user_barangay": "-",
            "detail_modal_id": f"postDetails{post.id}",
            "request_modal_id": f"{request_type}RequestUsers{post.id}",
            "pin_url": reverse("dogadoption_admin:toggle_post_pin", args=[post.id]),
            "pin_label": (
                "Unpin post from homepage spotlight"
                if getattr(post, "is_pinned", False)
                else "Pin post to homepage spotlight"
            ),
            "edit_url": reverse("dogadoption_admin:update_post", args=[post.id]),
            "toggle_url": reverse("dogadoption_admin:toggle_post_phase", args=[post.id]),
            "toggle_phase": toggle_phase,
            "toggle_label": toggle_label,
            "finalize_url": reverse("dogadoption_admin:finalize_post", args=[post.id]),
            "finalize_status": "reunited" if section_key == "claim" else "adopted",
            "finalize_label": "Mark as Redeemed" if section_key == "claim" else "Mark as Adopted",
            "finalize_icon": "fas fa-home" if section_key == "claim" else "fas fa-heart",
            "delete_url": reverse("dogadoption_admin:delete_post", args=[post.id]),
            "allow_toggle": allow_toggle,
        }

    def _paginate_status(posts, section_definition):
        paginator = Paginator(posts, rows_per_page)
        page_obj = paginator.get_page(request.GET.get(section_definition["page_param"], 1))
        items = [
            _build_post_item(
                post,
                section_definition["key"],
                section_definition["request_type"],
                section_definition["allow_toggle"],
            )
            for post in page_obj.object_list
        ]
        return page_obj, items

    active_post_candidates = list(
        base_qs.filter(status__in=active_statuses).order_by("-created_at", "-id")
    )
    active_appointment_dates = Post.attach_active_appointment_dates(active_post_candidates)

    open_posts_by_phase = {"claim": [], "adopt": []}
    for post in active_post_candidates:
        phase = post.current_phase()
        if phase in open_posts_by_phase:
            open_posts_by_phase[phase].append(post)

    for phase_key, posts in open_posts_by_phase.items():
        posts.sort(
            key=lambda post: (
                -int(bool(getattr(post, "is_pinned", False))),
                -int(getattr(post, "pinned_at", None).timestamp()) if getattr(post, "pinned_at", None) else 0,
                -int(getattr(post, "total_request_count", 0) or 0),
                -post.created_at.timestamp(),
                -post.id,
            )
        )

    section_post_map = {
        "claim": open_posts_by_phase["claim"],
        "adopt": open_posts_by_phase["adopt"],
    }

    history_total = len(_get_cached_post_history_ids())
    status_sections = []
    modal_posts_by_id = {}
    for section_definition in section_definitions:
        page_obj, items = _paginate_status(section_post_map[section_definition["key"]], section_definition)
        section_payload = {
            **section_definition,
            "count": page_obj.paginator.count,
            "items": items,
            "page_obj": page_obj,
            "page_links": [
                {
                    "number": page_number,
                    "qs": _build_page_qs(section_definition["page_param"], page_number),
                    "is_active": page_number == page_obj.number,
                }
                for page_number in page_obj.paginator.page_range
            ],
            "prev_qs": (
                _build_page_qs(section_definition["page_param"], page_obj.previous_page_number())
                if page_obj.has_previous()
                else ""
            ),
            "next_qs": (
                _build_page_qs(section_definition["page_param"], page_obj.next_page_number())
                if page_obj.has_next()
                else ""
            ),
        }
        status_sections.append(section_payload)
        for item in items:
            modal_posts_by_id.setdefault(item["post"].id, item)

    for item in modal_posts_by_id.values():
        item["claim_requests"] = []
        item["adopt_requests"] = []
        item["primary_image_url"] = ""
        item["finalized_user_name"] = "-"
        item["finalized_user_barangay"] = "-"

    paged_post_ids = list(modal_posts_by_id.keys())
    if paged_post_ids:
        paged_requests = list(
            PostRequest.objects.filter(post_id__in=paged_post_ids)
            .select_related("user", "post")
            .only(
                "id",
                "post_id",
                "user_id",
                "request_type",
                "status",
                "appointment_date",
                "scheduled_appointment_date",
                "created_at",
                "post__id",
                "post__created_at",
                "post__claim_days",
                "post__status",
                "user__id",
                "user__username",
                "user__first_name",
                "user__last_name",
            )
            .order_by("-created_at")
        )
        Post.attach_active_appointment_dates(
            [req.post for req in paged_requests if getattr(req, "post", None)],
            active_appointment_dates,
        )

        requests_by_post_id = defaultdict(lambda: {"claim": [], "adopt": []})
        request_user_ids = set()
        for req in paged_requests:
            request_user_ids.add(req.user_id)
            if req.request_type in {"claim", "adopt"}:
                requests_by_post_id[req.post_id][req.request_type].append(req)

        profile_image_by_user_id = {}
        profile_address_by_user_id = {}
        if request_user_ids:
            for profile in Profile.objects.filter(user_id__in=request_user_ids).only(
                "user_id",
                "address",
                "profile_image",
            ):
                profile_address_by_user_id[profile.user_id] = (profile.address or "").strip()
                image_url = _safe_media_url(getattr(profile, "profile_image", None))
                if image_url:
                    profile_image_by_user_id[profile.user_id] = image_url

        for req in paged_requests:
            display_name = f"{(req.user.first_name or '').strip()} {(req.user.last_name or '').strip()}".strip()
            if not display_name:
                display_name = req.user.username
            req.user_display_name = display_name
            req.user_initials = _owner_initials(display_name)
            req.user_profile_image_url = profile_image_by_user_id.get(req.user_id, "")

        primary_image_by_post_id = {}
        for image in PostImage.objects.filter(post_id__in=paged_post_ids).only("post_id", "image").order_by("id"):
            if image.post_id in primary_image_by_post_id:
                continue
            image_url = _safe_media_url(image.image)
            if image_url:
                primary_image_by_post_id[image.post_id] = image_url

        for post_id, item in modal_posts_by_id.items():
            req_bucket = requests_by_post_id.get(post_id)
            if req_bucket:
                item["claim_requests"] = req_bucket["claim"]
                item["adopt_requests"] = req_bucket["adopt"]
                accepted_req = None
                if item["post"].status == "reunited":
                    accepted_req = next((r for r in req_bucket["claim"] if r.status == "accepted"), None)
                elif item["post"].status == "adopted":
                    accepted_req = next((r for r in req_bucket["adopt"] if r.status == "accepted"), None)

                if accepted_req:
                    item["finalized_user_name"] = accepted_req.user_display_name
                    user_address = profile_address_by_user_id.get(accepted_req.user_id, "")
                    user_barangay = _extract_barangay_from_address(user_address)
                    if not user_barangay and user_address:
                        user_barangay = user_address.split(",")[0].strip()
                    item["finalized_user_barangay"] = user_barangay or "-"
            item["primary_image_url"] = primary_image_by_post_id.get(post_id, "")

    paged_all_posts = list(modal_posts_by_id.values())

    global_dates = active_appointment_dates

    user_post_pending_count = (
        UserAdoptionPost.objects.filter(status="pending_review").count()
        + MissingDogPost.objects.filter(status="pending_review").count()
    )

    return render(request, 'admin_home/post_list.html', {
        'all_posts': paged_all_posts,
        'status_sections': status_sections,
        'post_form': post_form,
        'show_create_modal': show_create_modal,
        'appointment_dates': [d.strftime('%Y-%m-%d') for d in global_dates],
        'show_appointment_modal': show_appointment_modal,
        'history_total': history_total,
        'return_to': request.get_full_path(),
        'user_post_pending_count': user_post_pending_count,
    })


@admin_required
def post_history(request):
    """Show expired unresolved posts in a dedicated, paginated history page."""
    history_context = _build_post_history_page(request, page_param="page", rows_per_page=10)
    page_obj = history_context["history_page_obj"]

    def _build_page_qs(page_num):
        params = request.GET.copy()
        params["page"] = str(page_num)
        return params.urlencode()

    def _build_filter_qs(filter_name):
        params = request.GET.copy()
        params["record_type"] = filter_name
        params["page"] = "1"
        return params.urlencode()

    return render(request, "admin_home/post_history.html", {
        **history_context,
        "history_prev_qs": _build_page_qs(page_obj.previous_page_number()) if page_obj.has_previous() else "",
        "history_next_qs": _build_page_qs(page_obj.next_page_number()) if page_obj.has_next() else "",
        "history_all_qs": _build_filter_qs("all"),
        "history_adopted_qs": _build_filter_qs("adopted"),
        "history_redeemed_qs": _build_filter_qs("redeemed"),
        "history_unresolved_qs": _build_filter_qs("unresolved"),
    })


@admin_required
@require_POST
def record_history_status(request, post_id):
    """Record an expired unresolved post as adopted or redeemed."""
    access = getattr(request, "admin_access", get_admin_access(request.user))
    if not access.get("can_create_posts"):
        messages.error(request, "You do not have permission to update history records.", extra_tags="post_list")
        return _redirect_to_safe_next(request, "dogadoption_admin:post_history")

    target_status = (request.POST.get("status") or "").strip().lower()
    if target_status not in {"adopted", "reunited"}:
        messages.error(request, "Invalid history status selected.", extra_tags="post_list")
        return _redirect_to_safe_next(request, "dogadoption_admin:post_history")

    with transaction.atomic():
        post = get_object_or_404(Post.objects.select_for_update(), id=post_id)
        post_name = post.display_title or post.caption or f"Post {post.id}"

        if post.status in {"reunited", "adopted"}:
            messages.warning(request, "This post has already been recorded as a final history entry.", extra_tags="post_list")
            return _redirect_to_safe_next(request, "dogadoption_admin:post_history")

        if not (post.is_history or post.is_expired()):
            messages.warning(
                request,
                "Only expired or archived unresolved posts can be recorded from history.",
                extra_tags="post_list",
            )
            return _redirect_to_safe_next(request, "dogadoption_admin:post_history")

        post.status = target_status
        post.phase_override = ""
        post.phase_override_started_at = None
        post.is_history = True
        post.is_pinned = False
        post.pinned_at = None
        post.save(
            update_fields=[
                "status",
                "phase_override",
                "phase_override_started_at",
                "is_history",
                "is_pinned",
                "pinned_at",
            ]
        )

        pending_requests = list(post.requests.filter(status="pending").only("id", "user_id"))
        if pending_requests:
            post.requests.filter(id__in=[req.id for req in pending_requests]).update(
                status="rejected",
                scheduled_appointment_date=None,
            )
            for pending_req in pending_requests:
                invalidate_user_notification_payload(pending_req.user_id)

    _touch_post_board_cache()
    messages.success(
        request,
        (
            f'Post "{post_name}" was recorded as redeemed in history.'
            if target_status == "reunited"
            else f'Post "{post_name}" was recorded as adopted in history.'
        ),
        extra_tags="post_list",
    )
    return _redirect_to_safe_next(request, "dogadoption_admin:post_history")


@admin_required
@require_http_methods(["GET", "POST"])
def appointment_calendar(request):
    """Maintain the shared appointment calendar used by post requests."""
    if request.method == 'POST':
        dates_raw = (request.POST.get('appointment_dates') or '').strip()
        if not _validate_and_save_global_appointment_dates(dates_raw, request.user):
            messages.error(request, "Past appointment dates are locked and cannot be changed.")
        else:
            messages.success(request, "Appointment dates saved.")
    global_dates = _get_active_global_appointment_dates()

    return render(request, 'admin_home/appointment_calendar.html', {
        'appointment_dates': [d.strftime('%Y-%m-%d') for d in global_dates],
    })

@admin_required
def claim_requests(request, post_id):
    """List claim requests tied to a single rescued post."""
    return _render_post_request_list(
        request,
        post_id,
        "claim",
        "admin_claim/claim_requests.html",
    )

@admin_required
def adoption_requests(request, post_id):
    """List adoption requests tied to a single rescued post."""
    return _render_post_request_list(
        request,
        post_id,
        "adopt",
        "admin_adoption/adoption_request.html",
    )

@admin_required
@require_POST
def update_request(request, req_id, action):
    """Accept or reject a claim/adoption request and update the related post."""
    action = (action or "").strip().lower()

    with transaction.atomic():
        req = get_object_or_404(
            PostRequest.objects.select_related("post").select_for_update(),
            id=req_id,
        )
        post = Post.objects.select_for_update().get(id=req.post_id)

        if action not in {'accept', 'reject'}:
            messages.error(request, "Unsupported action.")
            return _build_request_redirect_or_next(request, req)

        if req.status != 'pending':
            messages.warning(request, "This request has already been processed.")
            return _build_request_redirect_or_next(request, req)

        if action == 'accept':
            if not req.verification_ready:
                approval_open_at = req.approval_available_at
                if approval_open_at:
                    messages.error(
                        request,
                        "Approval opens after {}.".format(
                            timezone.localtime(approval_open_at).strftime("%b %d, %Y %I:%M %p")
                        ),
                    )
                else:
                    messages.error(request, "This request is still in its verification window.")
                return _build_request_redirect_or_next(request, req)

            scheduled_date_raw = (request.POST.get('scheduled_appointment_date') or '').strip()
            if scheduled_date_raw:
                scheduled_date = parse_date(scheduled_date_raw)
                if not scheduled_date:
                    messages.error(request, "Please select a valid appointment date.")
                    return _build_request_redirect_or_next(request, req)
                is_available = GlobalAppointmentDate.objects.filter(
                    appointment_date=scheduled_date,
                    is_active=True,
                ).exists()
                if not is_available:
                    messages.error(request, "Selected appointment date is not in the available schedule.")
                    return _build_request_redirect_or_next(request, req)
            else:
                scheduled_date = req.appointment_date

            req.status = 'accepted'
            req.scheduled_appointment_date = scheduled_date

            if req.request_type == 'claim':
                post.status = 'reunited'
            elif req.request_type == 'adopt':
                post.status = 'adopted'
            post.phase_override = ""
            post.phase_override_started_at = None
            post.save(update_fields=['status', 'phase_override', 'phase_override_started_at'])

            post.requests.filter(status='pending').exclude(id=req.id).update(
                status='rejected',
                scheduled_appointment_date=None,
            )
        else:
            req.status = 'rejected'
            req.scheduled_appointment_date = None

        reviewed_at = timezone.now()
        req.save(update_fields=['status', 'scheduled_appointment_date'])
        remember_request_reviewed_at(req.id, reviewed_at)
        invalidate_user_notification_payload(req.user_id)
        cache.delete(POST_HISTORY_CACHE_KEY)
        bump_user_home_feed_namespace()

    return _build_request_redirect_or_next(request, req)
# =============================================================================
# Navigation 2/5: Request
# Covers capture-request operations and supporting request review tools.
# =============================================================================

@admin_required
@require_http_methods(["GET", "POST"])
def admin_dog_capture_requests(request):
    """Manage incoming dog-capture requests and catcher contact details."""
    def _redirect_to_requests(default_tab="pending"):
        next_url = (request.POST.get("next") or "").strip()
        if next_url and url_has_allowed_host_and_scheme(
            next_url,
            allowed_hosts={request.get_host()},
            require_https=request.is_secure(),
        ):
            return redirect(next_url)

        return redirect(f"{reverse('dogadoption_admin:requests')}?tab={default_tab}")

    if request.method == 'POST':
        action = request.POST.get('action')

        if action == 'bulk_mark_captured':
            selected_ids = [
                int(value)
                for value in request.POST.getlist('selected_request_ids')
                if str(value).isdigit()
            ]
            scheduled_qs = DogCaptureRequest.objects.filter(
                id__in=selected_ids,
                status='accepted',
            )
            updated_count = scheduled_qs.count()
            if updated_count:
                scheduled_qs.update(
                    status='captured',
                    assigned_admin=request.user,
                    captured_at=timezone.now(),
                )
                messages.success(request, f"{updated_count} scheduled request(s) marked as done.")
            else:
                messages.warning(request, "Select at least one scheduled request to mark as done.")

            return _redirect_to_requests("accepted")

        elif action == 'reschedule_single':
            request_id = request.POST.get('request_id')
            req = DogCaptureRequest.objects.filter(
                id=request_id,
                status='accepted',
            ).select_related('requested_by', 'requested_by__profile').first()
            if not req:
                messages.error(request, "Scheduled request not found.")
                return _redirect_to_requests("accepted")

            scheduled_raw = (request.POST.get('scheduled_date') or '').strip()
            scheduled_date = parse_date(scheduled_raw) if scheduled_raw else None
            if not scheduled_date:
                messages.error(request, "Please select an available appointment date.")
                return _redirect_to_requests("accepted")

            is_available = GlobalAppointmentDate.objects.filter(
                appointment_date=scheduled_date,
                is_active=True,
            ).exists()
            if not is_available:
                messages.error(request, "Selected appointment date is not in the active admin schedule.")
                return _redirect_to_requests("accepted")

            scheduled_time = (
                timezone.localtime(req.scheduled_date).time().replace(second=0, microsecond=0)
                if req.scheduled_date
                else time(hour=9, minute=0)
            )
            req.scheduled_date = timezone.make_aware(
                datetime.combine(scheduled_date, scheduled_time),
                timezone.get_current_timezone(),
            )
            req.assigned_admin = request.user
            req.save(update_fields=['scheduled_date', 'assigned_admin'])

            messages.success(request, "Scheduled request updated.")
            return _redirect_to_requests("accepted")

        elif action == 'bulk_reschedule':
            selected_ids = [
                int(value)
                for value in request.POST.getlist('selected_request_ids')
                if str(value).isdigit()
            ]
            scheduled_raw = (request.POST.get('scheduled_date') or '').strip()
            scheduled_date = parse_date(scheduled_raw) if scheduled_raw else None
            if not selected_ids:
                messages.warning(request, "Select at least one scheduled request to update.")
                return _redirect_to_requests("accepted")
            if not scheduled_date:
                messages.error(request, "Please select an available appointment date.")
                return _redirect_to_requests("accepted")

            is_available = GlobalAppointmentDate.objects.filter(
                appointment_date=scheduled_date,
                is_active=True,
            ).exists()
            if not is_available:
                messages.error(request, "Selected appointment date is not in the active admin schedule.")
                return _redirect_to_requests("accepted")

            selected_requests = list(
                DogCaptureRequest.objects.filter(
                    id__in=selected_ids,
                    status='accepted',
                )
            )
            if not selected_requests:
                messages.warning(request, "Selected scheduled requests were not found.")
                return _redirect_to_requests("accepted")

            for req in selected_requests:
                scheduled_time = (
                    timezone.localtime(req.scheduled_date).time().replace(second=0, microsecond=0)
                    if req.scheduled_date
                    else time(hour=9, minute=0)
                )
                req.scheduled_date = timezone.make_aware(
                    datetime.combine(scheduled_date, scheduled_time),
                    timezone.get_current_timezone(),
                )
                req.assigned_admin = request.user
                req.save(update_fields=['scheduled_date', 'assigned_admin'])

            messages.success(request, f"{len(selected_requests)} scheduled request(s) updated.")
            return _redirect_to_requests("accepted")

        return _redirect_to_requests()

    rows_per_page = 10
    valid_tabs = {"pending", "accepted", "declined"}
    active_tab = (request.GET.get("tab") or "pending").strip().lower()
    if active_tab not in valid_tabs:
        active_tab = "pending"

    base_qs = _dog_capture_request_board_queryset().filter(request_type="surrender")
    status_totals = base_qs.aggregate(
        pending_total=Count("id", filter=Q(status="pending")),
        accepted_total=Count("id", filter=Q(status="accepted")),
        captured_total=Count("id", filter=Q(status="captured")),
        declined_total=Count("id", filter=Q(status="declined")),
    )

    def _paginate_status(status_key, page_param):
        filtered_qs = base_qs.filter(status=status_key)
        page_obj = Paginator(filtered_qs, rows_per_page).get_page(
            request.GET.get(page_param, 1)
        )
        items = list(page_obj.object_list)
        for req in items:
            _enrich_capture_request_display(req)
        return page_obj, items, page_obj.paginator.count

    pending_page_obj, pending_requests, pending_total = _paginate_status(
        "pending", "pending_page"
    )
    captured_page_obj, captured_requests, captured_total = _paginate_status(
        "captured", "captured_page"
    )
    declined_page_obj, declined_requests, declined_total = _paginate_status(
        "declined", "declined_page"
    )

    accepted_date_raw = (request.GET.get("accepted_date") or "").strip()
    accepted_date_filter = parse_date(accepted_date_raw) if accepted_date_raw else None
    accepted_total = int(status_totals.get("accepted_total") or 0)
    today = timezone.localdate()
    today_start = timezone.make_aware(
        datetime.combine(today, time.min),
        timezone.get_current_timezone(),
    )

    accepted_qs = (
        base_qs.filter(status='accepted')
        .annotate(
            future_first_flag=Case(
                When(scheduled_date__lt=today_start, then=Value(1)),
                default=Value(0),
                output_field=IntegerField(),
            ),
            no_schedule_last_flag=Case(
                When(scheduled_date__isnull=True, then=Value(1)),
                default=Value(0),
                output_field=IntegerField(),
            ),
            walk_in_last_flag=Case(
                When(submission_type='walk_in', then=Value(1)),
                default=Value(0),
                output_field=IntegerField(),
            ),
            sort_barangay=Lower(
                Trim(Coalesce("barangay", Value(""), output_field=CharField()))
            ),
            sort_location=Lower(
                Trim(
                    Coalesce(
                        "manual_full_address",
                        "barangay",
                        "city",
                        "requested_by__profile__address",
                        Value(""),
                        output_field=CharField(),
                    )
                )
            ),
        )
    )
    if accepted_date_filter:
        accepted_qs = accepted_qs.filter(scheduled_date__date=accepted_date_filter)

    accepted_qs = accepted_qs.order_by(
        "future_first_flag",
        "no_schedule_last_flag",
        "scheduled_date",
        "walk_in_last_flag",
        "sort_barangay",
        "sort_location",
        "created_at",
        "id",
    )
    accepted_page_obj = Paginator(accepted_qs, rows_per_page).get_page(
        request.GET.get("accepted_page", 1)
    )
    accepted_filtered_total = accepted_page_obj.paginator.count
    accepted_requests = list(accepted_page_obj.object_list)
    for req in accepted_requests:
        _enrich_capture_request_display(req)

    map_points_qs = list(
        _dog_capture_request_map_queryset().filter(
            request_type='surrender',
            status='pending',
            latitude__isnull=False,
            longitude__isnull=False,
        )[:400]
    )
    map_point_image_urls = _dog_capture_request_first_image_urls(
        [req.id for req in map_points_qs]
    )
    map_points = []
    for req in map_points_qs:
        _enrich_capture_request_display(req)
        map_points.append({
            'id': req.id,
            'user': req.requested_by.username,
            'requester_name': req.requester_full_name,
            'requester_phone': req.requester_phone,
            'requester_address': req.requester_address,
            'reason': req.get_reason_display(),
            'status': req.get_status_display(),
            'status_key': req.status,
            'request_type_key': req.request_type,
            'request_type_label': req.get_request_type_display(),
            'submission_type_key': req.submission_type or '',
            'submission_type_label': req.get_submission_type_display() if req.submission_type else '',
            'lat': float(req.latitude),
            'lng': float(req.longitude),
            'created_at': req.created_at.strftime('%b %d, %Y %I:%M %p'),
            'barangay': req.display_barangay,
            'location_label': req.location_label,
            'image_url': map_point_image_urls.get(req.id) or _safe_media_url(req.image),
        })

    available_appointment_dates = _get_available_appointment_dates()
    return render(request, 'admin_request/request.html', {
        'requests': bool(
            (status_totals.get("pending_total") or 0)
            or (status_totals.get("accepted_total") or 0)
            or (status_totals.get("captured_total") or 0)
            or (status_totals.get("declined_total") or 0)
        ),
        'pending_requests': pending_requests,
        'accepted_requests': accepted_requests,
        'captured_requests': captured_requests,
        'declined_requests': declined_requests,
        'pending_page_obj': pending_page_obj,
        'accepted_page_obj': accepted_page_obj,
        'captured_page_obj': captured_page_obj,
        'declined_page_obj': declined_page_obj,
        'pending_total': pending_total,
        'accepted_total': accepted_total,
        'captured_total': int(status_totals.get("captured_total") or 0),
        'declined_total': int(status_totals.get("declined_total") or 0),
        'accepted_filtered_total': accepted_filtered_total,
        'accepted_selected_date_iso': accepted_date_filter.isoformat() if accepted_date_filter else '',
        'accepted_selected_date_display': accepted_date_filter.strftime('%b %d, %Y') if accepted_date_filter else '',
        'accepted_date_qs': f"&accepted_date={accepted_date_filter.isoformat()}" if accepted_date_filter else '',
        'accepted_calendar_dates': [slot.appointment_date.strftime('%Y-%m-%d') for slot in available_appointment_dates],
        'active_tab': active_tab,
        'map_points': map_points,
        'requests_return_to': request.get_full_path(),
    })

@admin_required
def update_dog_capture_request(request, pk):
    """Review, schedule, or close a single dog-capture request."""
    req = get_object_or_404(
        DogCaptureRequest.objects.select_related('requested_by', 'requested_by__profile').prefetch_related('images', 'landmark_images'),
        pk=pk
    )
    _enrich_capture_request_user(req)

    if request.method == 'POST':
        action = request.POST.get('action')

        if req.status == 'captured' and action in {'accept', 'decline'}:
            messages.warning(request, "Captured records are closed and cannot be re-opened.")
            return redirect('dogadoption_admin:update_dog_capture_request', pk=req.id)

        if action == 'accept':
            scheduled_raw = (request.POST.get('scheduled_date') or '').strip()
            scheduled_date = parse_date(scheduled_raw) if scheduled_raw else None
            if not scheduled_date:
                messages.error(request, "Please select an available appointment date.")
                return redirect('dogadoption_admin:update_dog_capture_request', pk=req.id)

            is_available = GlobalAppointmentDate.objects.filter(
                appointment_date=scheduled_date,
                is_active=True,
            ).exists()
            if not is_available:
                messages.error(request, "Selected appointment date is not in the active admin schedule.")
                return redirect('dogadoption_admin:update_dog_capture_request', pk=req.id)

            scheduled_time = (
                timezone.localtime(req.scheduled_date).time().replace(second=0, microsecond=0)
                if req.scheduled_date
                else time(hour=9, minute=0)
            )
            scheduled_dt = timezone.make_aware(
                datetime.combine(scheduled_date, scheduled_time),
                timezone.get_current_timezone(),
            )
            req.status = 'accepted'
            req.assigned_admin = request.user
            req.scheduled_date = scheduled_dt
            req.admin_message = request.POST.get('admin_message')
            req.captured_at = None
            req.save()

            messages.success(request, "Request accepted and scheduled.")

        elif action == 'mark_captured':
            if req.status != 'accepted':
                messages.error(request, "Only scheduled requests can be marked as captured.")
                return redirect('dogadoption_admin:update_dog_capture_request', pk=req.id)

            admin_message = (request.POST.get('admin_message') or '').strip()
            req.status = 'captured'
            req.assigned_admin = request.user
            req.captured_at = timezone.now()
            if admin_message:
                req.admin_message = admin_message
            req.save()

            messages.success(request, "Request marked as captured.")

        elif action == 'decline':
            req.status = 'declined'
            req.admin_message = request.POST.get('admin_message')
            req.assigned_admin = request.user
            req.scheduled_date = None
            req.captured_at = None
            req.save()

            messages.warning(request, "Request declined.")

        return redirect('dogadoption_admin:requests')

    available_dates = _get_available_appointment_dates()
    return render(request, 'admin_request/update_request.html', {
        'req': req,
        'appointment_dates': [slot.appointment_date.strftime('%Y-%m-%d') for slot in available_dates],
        'requested_appointment_date_iso': req.preferred_appointment_date.strftime('%Y-%m-%d') if req.preferred_appointment_date else '',
        'scheduled_appointment_date_iso': timezone.localtime(req.scheduled_date).date().isoformat() if req.scheduled_date else '',
    })

# =============================================================================
# Navigation 4/5: Announcement
# Covers public admin announcements published to the user-facing app.
# =============================================================================

def _admin_announcement_queryset():
    return (
        DogAnnouncement.objects.select_related('created_by', 'created_by__profile')
        .only(
            'id',
            'title',
            'content',
            'category',
            'display_bucket',
            'background_image',
            'background_color',
            'created_at',
            'created_by__id',
            'created_by__username',
            'created_by__first_name',
            'created_by__last_name',
            'created_by__profile__profile_image',
        )
        .prefetch_related('images')
        .order_by('-created_at')
    )


@admin_required
def announcement_list(request):
    """Render the announcement feed shown in the admin announcement module."""
    announcement_page_obj = Paginator(
        _admin_announcement_queryset(),
        ADMIN_ANNOUNCEMENT_PAGE_SIZE,
    ).get_page(request.GET.get("page", 1))
    announcements = list(announcement_page_obj.object_list)
    default_admin_avatar_url = static("images/officialseal.webp")
    for post in announcements:
        profile = getattr(post.created_by, "profile", None)
        image_url = _safe_media_url(getattr(profile, "profile_image", None))
        post.admin_profile_image_url = image_url or default_admin_avatar_url

    return render(request, 'admin_announcement/announcement.html', {
        'announcements': announcements,
        'announcement_page_obj': announcement_page_obj,
        'category_options': ANNOUNCEMENT_CATEGORY_OPTIONS,
    })

@admin_required
def announcement_create(request):
    """Redirect announcement creation back to the category picker view."""
    return redirect("dogadoption_admin:admin_announcements")


@admin_required
@require_http_methods(["GET", "POST"])
def announcement_create_form(request, category_slug):
    """Create an announcement for the selected category."""
    category_option = ANNOUNCEMENT_CATEGORY_BY_SLUG.get(category_slug)
    if not category_option:
        raise Http404("Announcement category not found.")

    if request.method == "GET":
        return redirect("dogadoption_admin:admin_announcements")

    if request.method == "POST":
        title = (request.POST.get("title") or "").strip()
        content = (request.POST.get("content") or "").strip()
        background_color = (request.POST.get("background_color") or "#eeedf3").strip()
        uploaded_images = request.FILES.getlist("background_images")
        if not uploaded_images and request.FILES.get("background_image"):
            uploaded_images = [request.FILES.get("background_image")]
        schedule_raw = request.POST.get("schedule_data")
        schedule = None

        if schedule_raw:
            try:
                schedule = json.loads(schedule_raw)
            except json.JSONDecodeError:
                schedule = None

        if not content:
            messages.error(request, "Post content is required.")
            return redirect("dogadoption_admin:admin_announcements")

        primary_image = uploaded_images[0] if uploaded_images else None
        post = DogAnnouncement.objects.create(
            title=title or category_option["label"],
            content=content,
            category=category_option["value"],
            background_color=background_color,
            background_image=primary_image,
            schedule_data=schedule,
            created_by=request.user,
        )
        for image in uploaded_images[1:]:
            DogAnnouncementImage.objects.create(announcement=post, image=image)
        bump_user_home_feed_namespace()
        invalidate_user_notification_content()
        messages.success(request, f"{category_option['label']} post published.")

        return redirect("dogadoption_admin:admin_announcements")
    return redirect("dogadoption_admin:admin_announcements")


@admin_required
def announcement_edit(request, post_id):
    """Edit an existing announcement and optionally replace its images."""
    post = DogAnnouncement.objects.get(id=post_id)

    if request.method == "POST":
        post.title = (request.POST.get("title") or post.title).strip()
        post.content = (request.POST.get("content") or post.content).strip()
        category = request.POST.get("category", post.category)
        uploaded_images = request.FILES.getlist("background_images")
        if not uploaded_images and request.FILES.get("background_image"):
            uploaded_images = [request.FILES.get("background_image")]
        if category in ANNOUNCEMENT_CATEGORY_BY_VALUE:
            post.category = category
        post.background_color = (
            request.POST.get("background_color") or post.background_color
        )
        if uploaded_images:
            post.background_image = uploaded_images[0]
        post.save()
        if uploaded_images:
            post.images.all().delete()
            for image in uploaded_images[1:]:
                DogAnnouncementImage.objects.create(announcement=post, image=image)
        bump_user_home_feed_namespace()
        invalidate_user_notification_content()
        messages.success(request, "Announcement updated.")
        return redirect("dogadoption_admin:admin_announcements")

    return render(request, "admin_announcement/edit_announcement.html", {
        "post": post,
        "category_options": ANNOUNCEMENT_CATEGORY_OPTIONS,
    })

@admin_required
@require_POST
def announcement_update_bucket(request, post_id):
    """Change which announcement bucket the post belongs to."""
    post = get_object_or_404(DogAnnouncement.objects.only("id", "display_bucket"), id=post_id)
    bucket = (request.POST.get("bucket") or "").strip().lower()

    if bucket not in ANNOUNCEMENT_BUCKET_VALUES:
        return JsonResponse({
            "ok": False,
            "error": "Invalid announcement bucket.",
        }, status=400)

    if post.display_bucket != bucket:
        post.display_bucket = bucket
        post.save(update_fields=["display_bucket"])

    return JsonResponse({
        "ok": True,
        "bucket": post.display_bucket,
        "bucket_label": post.get_display_bucket_display(),
    })

@admin_required
@require_POST
def announcement_delete(request, post_id):
    """Delete an announcement and invalidate related user-facing content."""
    post = get_object_or_404(DogAnnouncement, id=post_id)
    post.delete()
    bump_user_home_feed_namespace()
    invalidate_user_notification_content()
    messages.success(request, "Announcement deleted.")

    return redirect("dogadoption_admin:admin_announcements")

# =============================================================================
# Shared admin utilities
# Covers pages that support admin work but are not part of the five sidebar links.
# =============================================================================

def _admin_user_management_queryset():
    """Build the base queryset used by the admin user-management screens."""
    return (
        User.objects.filter(is_staff=False)
        .select_related('profile', 'violation_summary', 'violation_summary__latest_notification')
        .annotate(
            claim_violation_count=Count(
                'postrequest',
                filter=Q(postrequest__request_type='claim'),
                distinct=True,
            ),
            citation_violation_count=Count('citation', distinct=True),
        )
        .annotate(
            calculated_violations=F('claim_violation_count') + F('citation_violation_count'),
            managed_violation_count=Coalesce('violation_summary__violation_count', Value(0)),
        )
        .annotate(
            effective_violation_count=Case(
                When(calculated_violations__gt=0, then=F('calculated_violations')),
                default=F('managed_violation_count'),
                output_field=IntegerField(),
            )
        )
    )


def _admin_user_display_name(user):
    full_name = " ".join(part for part in [user.first_name, user.last_name] if part).strip()
    return full_name or user.username


def _get_user_violation_summary_or_none(user):
    if user is None:
        return None
    try:
        return user.violation_summary
    except UserViolationSummary.DoesNotExist:
        return None


def _get_effective_violation_count(user):
    if user is None:
        return 0
    calculated_count = int(getattr(user, "calculated_violations", 0) or 0)
    managed_count = getattr(user, "managed_violation_count", None)
    if managed_count is None:
        summary = _get_user_violation_summary_or_none(user)
        managed_count = getattr(summary, "violation_count", 0) if summary else 0
    managed_count = int(managed_count or 0)
    return calculated_count if calculated_count > 0 else managed_count


def _build_violation_notification_status(summary, violation_count):
    latest_notification = getattr(summary, "latest_notification", None) if summary else None
    if latest_notification:
        if latest_notification.letter_status == UserViolationNotification.STATUS_PRINTED:
            return {"label": "Letter printed", "tone": "success"}
        return {"label": "Notice generated", "tone": "warning"}
    if violation_count >= VIOLATION_WARNING_THRESHOLD:
        return {"label": "Pending notice", "tone": "warning"}
    if violation_count > 0:
        return {"label": "Monitoring", "tone": "neutral"}
    return {"label": "No notice", "tone": "muted"}


def _build_admin_user_row_payloads(users):
    rows = []
    default_photo_url = static("images/default-user-image.jpg")
    for user in users:
        profile = _get_profile_or_none(user)
        summary = _get_user_violation_summary_or_none(user)
        violation_count = _get_effective_violation_count(user)
        rows.append(
            {
                "id": user.id,
                "full_name": _admin_user_display_name(user),
                "username": user.username,
                "date_joined": user.date_joined,
                "photo_url": _safe_media_url(getattr(profile, "profile_image", None)) or default_photo_url,
                "violation_count": violation_count,
                "notification_status": _build_violation_notification_status(summary, violation_count),
                "profile_url": reverse("dogadoption_admin:registration_owner_profile", args=[user.id]),
                "violation_url": reverse("dogadoption_admin:admin_user_violations", args=[user.id]),
            }
        )
    return rows


def _build_violation_threshold_message(user, violation_count):
    return (
        f"{_admin_user_display_name(user)} has reached {violation_count} recorded violations and may be "
        f"subject to disciplinary action based on system policy."
    )


def _ensure_violation_threshold_notification(user, summary, violation_count):
    threshold_count = int(violation_count or 0)
    if summary is not None:
        threshold_count = max(threshold_count, int(getattr(summary, "violation_count", 0) or 0))
    if threshold_count < VIOLATION_WARNING_THRESHOLD:
        return None, False

    if summary is None:
        summary = UserViolationSummary.objects.create(user=user, violation_count=threshold_count)
    elif threshold_count > int(getattr(summary, "violation_count", 0) or 0):
        summary.violation_count = threshold_count
        summary.save(update_fields=["violation_count", "updated_at"])

    notification, created = UserViolationNotification.objects.get_or_create(
        summary=summary,
        trigger_violation_count=VIOLATION_WARNING_THRESHOLD,
        defaults={
            "title": "Notice of Final Warning",
            "message": _build_violation_threshold_message(user, threshold_count),
        },
    )

    event_key = f"user-violation-threshold:{user.id}:{VIOLATION_WARNING_THRESHOLD}"
    admin_notification = notification.admin_notification
    if admin_notification is None:
        admin_notification = AdminNotification.objects.filter(event_key=event_key).first()
        if admin_notification is None:
            admin_notification = AdminNotification.objects.create(
                title=f"User reached {VIOLATION_WARNING_THRESHOLD} violations",
                message=notification.message,
                url=reverse("dogadoption_admin:admin_user_violations", args=[user.id]),
                event_key=event_key,
            )
            cache.delete(ADMIN_NOTIFICATIONS_CACHE_KEY)
        notification.admin_notification = admin_notification
        notification.save(update_fields=["admin_notification"])

    if summary.latest_notification_id != notification.id:
        summary.latest_notification = notification
        summary.save()

    return notification, created


def _build_violation_letter_context(user, summary, latest_notification=None, violation_count=None):
    profile = _get_profile_or_none(user)
    address = (getattr(profile, "address", "") or "").strip()
    dogs = list(
        Dog.objects.filter(owner_user=user)
        .only("id", "name", "barangay", "date_registered")
        .order_by("-date_registered", "-id")[:4]
    )
    registration_ids = [str(dog.id) for dog in dogs]
    registration_reference = ", ".join(registration_ids[:3]) if registration_ids else "No linked registration record"
    if len(registration_ids) > 3:
        registration_reference = f"{registration_reference} +{len(registration_ids) - 3} more"

    barangay = next(
        ((dog.barangay or "").strip() for dog in dogs if (dog.barangay or "").strip()),
        "",
    ) or _extract_barangay_from_address(address)

    violation_count = int(
        violation_count
        if violation_count is not None
        else (getattr(summary, "violation_count", 0) if summary else 0)
    )
    is_final_warning = violation_count >= VIOLATION_WARNING_THRESHOLD
    title = (
        latest_notification.title
        if latest_notification
        else ("Notice of Final Warning" if is_final_warning else "Notice of Violation Warning")
    )
    message = (
        latest_notification.message
        if latest_notification and latest_notification.message
        else (
            f"This notice is issued to inform you that you currently have {violation_count} recorded "
            f"violation{'s' if violation_count != 1 else ''} in the registration system."
        )
    )
    follow_up = (
        "This serves as a final warning. Further non-compliance may result in disciplinary action "
        "based on system policy."
        if is_final_warning
        else "Please comply with all registration and system policies to avoid additional sanctions."
    )

    return {
        "title": title,
        "document_date": timezone.localdate(),
        "display_name": _admin_user_display_name(user),
        "user_id": user.id,
        "registration_reference": registration_reference,
        "violation_count": violation_count,
        "barangay": barangay or "-",
        "address": address or "-",
        "message": message,
        "follow_up": follow_up,
        "is_final_warning": is_final_warning,
        "status": getattr(latest_notification, "letter_status", "draft"),
        "printed_at": getattr(latest_notification, "printed_at", None),
        "office_name": VIOLATION_OFFICE_NAME,
        "office_address_lines": VIOLATION_OFFICE_ADDRESS_LINES,
        "signatory_name": VIOLATION_SIGNATORY_NAME,
        "signatory_role": VIOLATION_SIGNATORY_ROLE,
    }


def _paginated_admin_user_directory_context(request, users_qs):
    """Slice the queryset to one page of rows and build list payloads (avoids loading the full directory)."""
    paginator = Paginator(users_qs, ADMIN_USERS_PER_PAGE)
    page_obj = paginator.get_page(request.GET.get("page", 1))
    page_users = list(page_obj.object_list)
    return {
        "users": _build_admin_user_row_payloads(page_users),
        "user_count": paginator.count,
        "page_obj": page_obj,
    }


@admin_required
def admin_users(request):
    """List non-staff users together with their violation counts."""
    query = " ".join((request.GET.get('q') or '').split()).strip()

    users = _admin_user_management_queryset()

    if query:
        users = users.filter(
            Q(first_name__icontains=query) |
            Q(last_name__icontains=query) |
            Q(username__icontains=query)
        )

    users = users.order_by('-effective_violation_count', 'first_name', 'last_name', 'username')
    context = _paginated_admin_user_directory_context(request, users)
    context["query"] = query
    context["is_search_view"] = False
    return render(request, 'admin_user/users.html', context)


@admin_required
def admin_user_detail(request, id):
    """Show the full admin-side detail page for a selected user."""
    user = get_object_or_404(
        _admin_user_management_queryset().prefetch_related('faceimage_set'),
        id=id
    )
    return render(request, 'admin_user/user_detail.html', {'user': user})


@admin_required
def admin_user_search_results(request):
    """Render the filtered user-management page."""
    query = " ".join((request.GET.get('q') or '').split()).strip()

    results = _admin_user_management_queryset().filter(
        Q(first_name__icontains=query) |
        Q(last_name__icontains=query) |
        Q(username__icontains=query)
    ).order_by('-effective_violation_count', 'first_name', 'last_name', 'username')

    context = _paginated_admin_user_directory_context(request, results)
    context["query"] = query
    context["is_search_view"] = True
    return render(request, 'admin_user/users.html', context)


@admin_required
def admin_user_violations(request, id):
    """Show the admin violation status and letter preview for a user."""
    user = get_object_or_404(_admin_user_management_queryset(), id=id)
    summary = _get_user_violation_summary_or_none(user)
    violation_count = _get_effective_violation_count(user)
    latest_notification = getattr(summary, "latest_notification", None) if summary else None
    if summary:
        if latest_notification is None:
            latest_notification = summary.notifications.order_by("-created_at", "-id").first()
    if violation_count >= VIOLATION_WARNING_THRESHOLD or (
        summary and int(getattr(summary, "violation_count", 0) or 0) >= VIOLATION_WARNING_THRESHOLD
    ):
        latest_notification, _ = _ensure_violation_threshold_notification(user, summary, violation_count)

    context = {
        "managed_user": user,
        "managed_profile": _get_profile_or_none(user),
        "managed_violation_count": violation_count,
        "latest_notification": latest_notification,
        "letter": _build_violation_letter_context(user, summary, latest_notification, violation_count=violation_count),
        "print_url": reverse("dogadoption_admin:admin_user_violation_letter", args=[user.id]),
        "profile_url": reverse("dogadoption_admin:registration_owner_profile", args=[user.id]),
    }
    return render(request, "admin_user/violation_detail.html", context)


@admin_required
def admin_user_violation_letter(request, id):
    """Render the printable violation letter for a user."""
    user = get_object_or_404(_admin_user_management_queryset(), id=id)
    summary = _get_user_violation_summary_or_none(user)
    violation_count = _get_effective_violation_count(user)
    if violation_count <= 0:
        messages.warning(request, "There are no recorded violations to print for this user.")
        return redirect("dogadoption_admin:admin_user_violations", id=user.id)

    latest_notification = getattr(summary, "latest_notification", None)
    if violation_count >= VIOLATION_WARNING_THRESHOLD or (
        summary and int(getattr(summary, "violation_count", 0) or 0) >= VIOLATION_WARNING_THRESHOLD
    ):
        latest_notification, _ = _ensure_violation_threshold_notification(user, summary, violation_count)
        if latest_notification and latest_notification.letter_status != UserViolationNotification.STATUS_PRINTED:
            latest_notification.letter_status = UserViolationNotification.STATUS_PRINTED
            latest_notification.printed_at = timezone.now()
            latest_notification.save(update_fields=["letter_status", "printed_at"])

    context = {
        "managed_user": user,
        "letter": _build_violation_letter_context(user, summary, latest_notification, violation_count=violation_count),
        "back_url": reverse("dogadoption_admin:admin_user_violations", args=[user.id]),
    }
    return render(request, "admin_user/violation_letter_print.html", context)


def _build_staff_permission_groups(form):
    groups = []
    for group in STAFF_PERMISSION_GROUPS:
        groups.append(
            {
                "title": group["title"],
                "description": group.get("description", ""),
                "note": group.get("note", ""),
                "items": [
                    {
                        "field": form[item["name"]],
                        "label": item["label"],
                        "help": item["help"],
                    }
                    for item in group["items"]
                ],
            }
        )
    return groups


def _build_staff_management_rows(bound_update_form=None, active_staff_panel=""):
    rows = []
    access_rows = StaffAccess.objects.select_related("user").order_by("user__username")
    for access_row in access_rows:
        staff_user = access_row.user
        if bound_update_form is not None and getattr(bound_update_form.instance, "pk", None) == staff_user.pk:
            form = bound_update_form
        else:
            form = ManagedStaffAccountForm(
                instance=staff_user,
                require_password=False,
                prefix=f"staff-{staff_user.pk}",
            )
        rows.append(
            {
                "user": staff_user,
                "form": form,
                "permission_groups": _build_staff_permission_groups(form),
                "permission_summary": get_staff_permission_summary(access_row),
                "panel_id": f"staff-{staff_user.pk}",
                "is_open": active_staff_panel == f"staff-{staff_user.pk}",
            }
        )
    return rows


@admin_required
def admin_edit_profile(request):
    """Allow the current admin to update login settings and manage staff accounts."""
    user = request.user
    profile, created = Profile.objects.get_or_create(
        user=user,
        defaults={
            "address": "",
            "age": 18,
            "consent_given": True
        }
    )
    access = getattr(request, "admin_access", get_admin_access(user))
    staff_create_form = ManagedStaffAccountForm(prefix="create-staff")
    bound_update_form = None
    active_staff_panel = ""

    if request.method == "POST":
        action = (request.POST.get("action") or "update_profile").strip()

        if action == "update_profile":
            username = (request.POST.get("username") or "").strip()
            current_password = request.POST.get("current_password") or ""
            password = request.POST.get("password") or ""
            confirm_password = request.POST.get("confirm_password") or ""

            has_error = False

            if not username:
                messages.error(request, "Username is required.")
                has_error = True
            elif User.objects.exclude(pk=user.pk).filter(username__iexact=username).exists():
                messages.error(request, "That username is already in use.")
                has_error = True

            if password or confirm_password:
                if not current_password:
                    messages.error(request, "Enter your current password first.")
                    has_error = True
                elif not user.check_password(current_password):
                    messages.error(request, "Current password is incorrect.")
                    has_error = True
                elif password != confirm_password:
                    messages.error(request, "Password confirmation does not match.")
                    has_error = True
                elif len(password) < 8:
                    messages.error(request, "Password must be at least 8 characters.")
                    has_error = True

            if not has_error:
                user.username = username
                if password:
                    user.set_password(password)
                user.save()
                if password:
                    update_session_auth_hash(request, user)
                messages.success(request, "Admin profile updated successfully.")
                return redirect("dogadoption_admin:admin_edit_profile")

        elif action == "create_staff":
            if not access.get("can_manage_staff_accounts"):
                messages.error(request, "Only the admin can manage staff accounts.")
                return redirect(access.get("landing_url") or reverse("dogadoption_admin:admin_edit_profile"))

            staff_create_form = ManagedStaffAccountForm(
                request.POST,
                prefix="create-staff",
            )
            active_staff_panel = "create-staff"
            if staff_create_form.is_valid():
                with transaction.atomic():
                    staff_user = staff_create_form.save()
                    Profile.objects.get_or_create(
                        user=staff_user,
                        defaults={
                            "address": "",
                            "age": 18,
                            "consent_given": True,
                        },
                    )
                messages.success(request, f"Staff account @{staff_user.username} created successfully.")
                return redirect("dogadoption_admin:admin_edit_profile")

        elif action == "update_staff":
            if not access.get("can_manage_staff_accounts"):
                messages.error(request, "Only the admin can manage staff accounts.")
                return redirect(access.get("landing_url") or reverse("dogadoption_admin:admin_edit_profile"))

            staff_user = get_object_or_404(
                User.objects.filter(is_staff=True, staff_access__isnull=False),
                pk=request.POST.get("staff_user_id"),
            )
            bound_update_form = ManagedStaffAccountForm(
                request.POST,
                instance=staff_user,
                require_password=False,
                prefix=f"staff-{staff_user.pk}",
            )
            active_staff_panel = f"staff-{staff_user.pk}"
            if bound_update_form.is_valid():
                with transaction.atomic():
                    updated_user = bound_update_form.save()
                    Profile.objects.get_or_create(
                        user=updated_user,
                        defaults={
                            "address": "",
                            "age": 18,
                            "consent_given": True,
                        },
                    )
                messages.success(request, f"Staff account @{updated_user.username} updated successfully.")
                return redirect("dogadoption_admin:admin_edit_profile")

        elif action == "toggle_staff":
            if not access.get("can_manage_staff_accounts"):
                messages.error(request, "Only the admin can manage staff accounts.")
                return redirect(access.get("landing_url") or reverse("dogadoption_admin:admin_edit_profile"))

            staff_user = get_object_or_404(
                User.objects.filter(is_staff=True, staff_access__isnull=False),
                pk=request.POST.get("staff_user_id"),
            )
            staff_user.is_active = not staff_user.is_active
            staff_user.save(update_fields=["is_active"])
            state_label = "activated" if staff_user.is_active else "deactivated"
            messages.success(request, f"Staff account @{staff_user.username} {state_label}.")
            return redirect("dogadoption_admin:admin_edit_profile")

        elif action == "delete_staff":
            if not access.get("can_manage_staff_accounts"):
                messages.error(request, "Only the admin can manage staff accounts.")
                return redirect(access.get("landing_url") or reverse("dogadoption_admin:admin_edit_profile"))

            staff_user = get_object_or_404(
                User.objects.filter(is_staff=True, staff_access__isnull=False),
                pk=request.POST.get("staff_user_id"),
            )
            deleted_username = staff_user.username
            with transaction.atomic():
                staff_user.delete()
            messages.success(request, f"Staff account @{deleted_username} deleted successfully.")
            return redirect("dogadoption_admin:admin_edit_profile")

    return render(
        request,
        "admin_profile/edit_profile.html",
        {
            "profile": profile,
            "staff_create_form": staff_create_form,
            "staff_create_permission_groups": _build_staff_permission_groups(staff_create_form),
            "staff_rows": _build_staff_management_rows(bound_update_form, active_staff_panel),
            "can_manage_staff_accounts": access.get("can_manage_staff_accounts"),
            "active_staff_panel": active_staff_panel,
        },
    )


@admin_required
def admin_notifications(request):
    """Display admin notifications and support marking all as read."""
    if request.method == "POST":
        action = request.POST.get("action")
        if action == "mark_all_read":
            AdminNotification.objects.filter(is_read=False).update(is_read=True)
            cache.delete(ADMIN_NOTIFICATIONS_CACHE_KEY)
        return redirect("dogadoption_admin:admin_notifications")

    notifications = AdminNotification.objects.all()
    return render(request, "admin_notifications/notifications.html", {
        "notifications": notifications,
    })


@admin_required
def notification_summary(request):
    """Return the current admin notification badge and dropdown data."""
    payload = _build_admin_notification_summary()
    notifications = []
    for item in payload.get("admin_latest_notifications", []):
        notifications.append({
            "id": item["id"],
            "title": item["title"],
            "message": item["message"],
            "created_label": timezone.localtime(item["created_at"]).strftime("%b %d, %Y %I:%M %p"),
            "is_read": item["is_read"],
            "read_url": reverse("dogadoption_admin:notification_read", args=[item["id"]]),
        })
    return JsonResponse({
        "unread_count": payload.get("admin_unread_notifications", 0),
        "notifications": notifications,
    })


@admin_required
@require_POST
def mark_notification_read(request, pk):
    """Mark one admin notification as read and follow its target link."""
    notif = get_object_or_404(AdminNotification, pk=pk)
    if not notif.is_read:
        notif.is_read = True
        notif.save(update_fields=["is_read"])
        cache.delete(ADMIN_NOTIFICATIONS_CACHE_KEY)
    target = notif.url or "dogadoption_admin:admin_notifications"
    return redirect(target)


@admin_required
def citation_print_lookup(request):
    """Resolve a numeric citation reference into its signed print URL."""
    raw_citation_id = (request.GET.get("citation_id") or "").strip()
    try:
        citation_id = int(raw_citation_id)
    except (TypeError, ValueError):
        messages.error(request, "Enter a valid citation ID.")
        return redirect("dogadoption_admin:citation_create")

    if citation_id < 1:
        messages.error(request, "Enter a valid citation ID.")
        return redirect("dogadoption_admin:citation_create")

    citation_exists = Citation.objects.filter(pk=citation_id).exists()
    if not citation_exists:
        messages.error(request, "Citation not found.")
        return redirect("dogadoption_admin:citation_create")

    return redirect("dogadoption_admin:citation_print", pk=citation_id)

# =============================================================================
# Navigation 5/5: Analytics
# Covers the analytics dashboard linked from the admin sidebar.
# =============================================================================


def _build_choice_count_chart(rows, choices):
    totals = {row["status"]: row["total"] for row in rows}
    return {
        "labels": [label for _, label in choices],
        "data": [totals.get(key, 0) for key, _ in choices],
    }


def _build_request_status_chart():
    request_matrix = {}
    for row in PostRequest.objects.values("request_type", "status").annotate(total=Count("id")):
        request_matrix.setdefault(row["request_type"], {})[row["status"]] = row["total"]

    request_status_display = {
        "pending": "Pending",
        "accepted": "Accepted",
        "rejected": "Rejected",
    }
    request_types = [key for key, _ in PostRequest.REQUEST_TYPE_CHOICES]
    return {
        "labels": [label for _, label in PostRequest.REQUEST_TYPE_CHOICES],
        "datasets": [
            {
                "label": request_status_display.get(status, status.title()),
                "data": [request_matrix.get(request_type, {}).get(status, 0) for request_type in request_types],
            }
            for status, _ in PostRequest.STATUS_CHOICES
        ],
    }


def _build_adoption_claim_trend_chart():
    rows = []
    years = set()
    trend_rows = (
        PostRequest.objects.filter(
            status="accepted",
            request_type__in=["claim", "adopt"],
            post__status__in=["reunited", "adopted"],
        )
        .annotate(activity_date=Coalesce("scheduled_appointment_date", TruncDate("created_at")))
        .exclude(activity_date__isnull=True)
        .values("request_type", "activity_date")
        .annotate(total=Count("id"))
        .order_by("activity_date", "request_type")
    )
    for row in trend_rows:
        activity_date = row["activity_date"]
        rows.append({
            "status": "claimed" if row["request_type"] == "claim" else "adopted",
            "date": activity_date.isoformat(),
            "total": row["total"],
        })
        years.add(activity_date.year)

    return {
        "rows": rows,
        "years": sorted(years),
    }


def _build_vaccination_breed_chart():
    vaccination_breed_counts = defaultdict(int)
    vaccination_breed_labels = {}
    vaccination_breed_years = set()
    vaccination_breed_trends = (
        VaccinationRecord.objects.exclude(registration__isnull=True)
        .exclude(registration__breed__isnull=True)
        .exclude(registration__breed__exact="")
        .values("date", "registration__breed")
        .annotate(total=Count("registration_id", distinct=True))
        .order_by("date", "registration__breed")
    )
    for row in vaccination_breed_trends:
        vaccination_date = row["date"]
        breed_raw = row["registration__breed"]
        breed_key = _normalize_breed_key(breed_raw)
        if not vaccination_date or not breed_key or _exclude_breed_from_chart(breed_raw):
            continue

        breed_type = _classify_breed_type(breed_raw)
        label_key = (breed_key, breed_type)
        vaccination_breed_labels.setdefault(label_key, _format_breed_label(breed_raw))
        vaccination_breed_counts[(vaccination_date, breed_key, breed_type)] += row["total"]
        vaccination_breed_years.add(vaccination_date.year)

    rows = []
    for (vaccination_date, breed_key, breed_type), total in sorted(
        vaccination_breed_counts.items(),
        key=lambda item: (item[0][0], item[0][1], item[0][2]),
    ):
        rows.append({
            "date": vaccination_date.isoformat(),
            "breed": vaccination_breed_labels[(breed_key, breed_type)],
            "animal_type": breed_type,
            "total": total,
        })

    return {
        "rows": rows,
        "years": sorted(vaccination_breed_years),
    }


def _build_rescue_barangay_trend_chart():
    events = []
    years = set()
    rescue_rows = (
        Post.objects.exclude(location__isnull=True)
        .exclude(location__exact="")
        .annotate(activity_date=Coalesce("rescued_date", TruncDate("created_at")))
        .exclude(activity_date__isnull=True)
        .values("location", "activity_date")
        .order_by("activity_date", "location")
    )
    for row in rescue_rows:
        location = _clean_barangay(row["location"])
        if not location:
            continue

        activity_date = row["activity_date"]
        if not activity_date:
            continue

        barangay_name = (
            _extract_barangay_from_address(location)
            or _resolve_barangay_name(location)
            or location
        )
        events.append({
            "barangay": barangay_name,
            "date": activity_date.isoformat(),
        })
        years.add(activity_date.year)

    return {
        "events": events,
        "years": sorted(years),
    }


def _build_vaccination_barangay_chart(today):
    events = []
    years = set()
    for record in (
        VaccinationRecord.objects.exclude(registration__isnull=True)
        .values(
            "registration_id",
            "registration__address",
            "date",
            "vaccine_expiry_date",
            "vaccination_expiry_date",
        )
    ):
        if record["date"]:
            years.add(record["date"].year)
        if record["vaccine_expiry_date"]:
            years.add(record["vaccine_expiry_date"].year)
        if record["vaccination_expiry_date"]:
            years.add(record["vaccination_expiry_date"].year)

        events.append({
            "registration_id": record["registration_id"],
            "barangay": _extract_barangay_from_address(record["registration__address"]) or "Unknown",
            "vaccination_date": record["date"].isoformat() if record["date"] else "",
            "vaccine_expiry_date": record["vaccine_expiry_date"].isoformat() if record["vaccine_expiry_date"] else "",
            "dog_vaccination_expiry_date": (
                record["vaccination_expiry_date"].isoformat()
                if record["vaccination_expiry_date"] else ""
            ),
        })

    return {
        "events": events,
        "years": sorted(years),
        "today": today.isoformat(),
    }


def _build_registered_barangay_chart():
    events = []
    years = set()
    for row in (
        Dog.objects.exclude(barangay__isnull=True)
        .exclude(barangay__exact="")
        .exclude(date_registered__isnull=True)
        .values("barangay", "date_registered")
        .order_by("date_registered", "barangay")
    ):
        registration_date = row["date_registered"]
        barangay_name = (row["barangay"] or "").strip()
        if not registration_date or not barangay_name:
            continue
        events.append({
            "barangay": barangay_name,
            "date": registration_date.isoformat(),
        })
        years.add(registration_date.year)

    return {
        "events": events,
        "years": sorted(years),
    }


def _build_analytics_dashboard_context():
    today = timezone.localdate()
    vaccinated_registration_ids = VaccinationRecord.objects.exclude(registration__isnull=True)
    return {
        "registered_owners": (
            DogRegistration.objects.exclude(owner_name__isnull=True)
            .exclude(owner_name__exact="")
            .annotate(owner_name_normalized=Lower(Trim("owner_name")))
            .values("owner_name_normalized")
            .distinct()
            .count()
        ),
        "adopted_dogs": Post.objects.filter(status="adopted").count(),
        "claimed_dogs": Post.objects.filter(status="reunited").count(),
        "vaccinated_dogs": vaccinated_registration_ids.values("registration_id").distinct().count(),
        "expired_vaccinations": (
            vaccinated_registration_ids.filter(
                Q(vaccine_expiry_date__lt=today) | Q(vaccination_expiry_date__lt=today)
            )
            .values("registration_id")
            .distinct()
            .count()
        ),
        "total_users": User.objects.filter(is_staff=False).count(),
        "total_posts": Post.objects.count(),
        "total_capture_requests": DogCaptureRequest.objects.count(),
        "total_registrations": DogRegistration.objects.count(),
        "post_status_chart": _build_choice_count_chart(
            Post.objects.values("status").annotate(total=Count("id")),
            Post.STATUS_CHOICES,
        ),
        "request_status_chart": _build_request_status_chart(),
        "capture_status_chart": _build_choice_count_chart(
            DogCaptureRequest.objects.values("status").annotate(total=Count("id")),
            DogCaptureRequest.STATUS_CHOICES,
        ),
        "vaccination_breed_chart": _build_vaccination_breed_chart(),
        "adoption_claim_trend_chart": _build_adoption_claim_trend_chart(),
        "rescue_barangay_trend_chart": _build_rescue_barangay_trend_chart(),
        "vaccination_barangay_chart": _build_vaccination_barangay_chart(today),
        "barangay_chart": _build_registered_barangay_chart(),
    }


@admin_required
def analytics_dashboard(request):
    """Build and cache the admin analytics dashboard context."""
    cached_context = cache.get(ANALYTICS_DASHBOARD_CACHE_KEY)
    if cached_context is not None:
        return render(request, "admin_analytics/dashboard.html", cached_context)

    context = _build_analytics_dashboard_context()
    cache.set(
        ANALYTICS_DASHBOARD_CACHE_KEY,
        context,
        ANALYTICS_DASHBOARD_CACHE_TTL_SECONDS,
    )
    return render(request, "admin_analytics/dashboard.html", context)

# =============================================================================
# Navigation 3/5: Register
# Covers registration, vaccination records, certificate exports, and citations.
# =============================================================================

# ---------------------------------------------------------------------------
# Register link 1/5: Registration
# ---------------------------------------------------------------------------
@admin_required
def register_dogs(request):
    """Register a new dog record from the admin registration form."""
    def _registration_error(text):
        messages.error(request, text, extra_tags="registration")

    def _registration_success(text):
        messages.success(request, text, extra_tags="registration")

    # Admin-controlled barangay and date stored in session
    barangay = request.session.get('barangay', '')
    date = request.session.get('date', '')

    if request.method == 'POST':
        barangay_input = request.POST.get('barangay', barangay)
        barangay = _resolve_barangay_name(barangay_input)
        date = request.POST.get('date', date)
        request.session['barangay'] = barangay
        request.session['date'] = date

        name = request.POST.get('name', '').strip()
        species = request.POST.get('species', 'Canine').strip()
        sex = request.POST.get('sex', 'M')
        age_value_raw = (request.POST.get('age_value') or '').strip()
        age_unit = (request.POST.get('age_unit') or 'years').strip()
        neutering = request.POST.get('neutering', 'No')
        color = request.POST.get('color', '').strip()
        owner_first_name = request.POST.get('owner_first_name', '').strip()
        owner_last_name = request.POST.get('owner_last_name', '').strip()
        owner_user_id = (request.POST.get("owner_user_id") or "").strip()
        uploaded_gallery_images = [img for img in request.FILES.getlist("dog_images") if img]
        uploaded_camera_images = [img for img in request.FILES.getlist("dog_camera_images") if img]
        uploaded_desktop_camera_images = [
            img for img in request.FILES.getlist("captured_camera_images") if img
        ]
        uploaded_images = (
            uploaded_gallery_images
            + uploaded_camera_images
            + uploaded_desktop_camera_images
        )

        if not barangay:
            _registration_error("Please select a valid barangay from the suggestions.")
            return redirect('dogadoption_admin:register_dogs')

        if not name or not owner_first_name or not owner_last_name:
            _registration_error("Dog Name and Owner First/Last Name are required.")
            return redirect('dogadoption_admin:register_dogs')

        if (owner_first_name or owner_last_name) and (not owner_first_name or not owner_last_name):
            _registration_error("Please provide both owner first name and last name.")
            return redirect('dogadoption_admin:register_dogs')

        owner_name, owner_name_key, owner_user = _resolve_registration_owner_identity(
            owner_first_name,
            owner_last_name,
            owner_user_id=owner_user_id,
        )
        if not owner_name or not owner_name_key:
            _registration_error("Please provide a valid owner first name and last name.")
            return redirect('dogadoption_admin:register_dogs')

        if species not in {"Canine", "Feline"}:
            _registration_error("Please select a valid species (Canine or Feline).")
            return redirect('dogadoption_admin:register_dogs')

        try:
            age_value = int(age_value_raw)
            if age_value <= 0:
                raise ValueError
        except (TypeError, ValueError):
            _registration_error("Age must be a valid positive number.")
            return redirect('dogadoption_admin:register_dogs')

        if age_unit not in {"months", "years"}:
            _registration_error("Please select a valid age unit.")
            return redirect('dogadoption_admin:register_dogs')

        age = f"{age_value} {'mos' if age_unit == 'months' else 'yrs'}"

        image_error = _validate_registration_images(uploaded_images)
        if image_error:
            _registration_error(image_error)
            return redirect('dogadoption_admin:register_dogs')

        try:
            date_registered = datetime.strptime(date, '%Y-%m-%d').date()
        except ValueError:
            _registration_error("Invalid date format.")
            return redirect('dogadoption_admin:register_dogs')
        
        formatted_address = f"{barangay}, Bayawan City, Negros Oriental"

        with transaction.atomic():
            owner_limit_query = _build_owner_limit_query(
                owner_name_key=owner_name_key,
                owner_name=owner_name,
                owner_user=owner_user,
            )
            owner_registered_count = (
                Dog.objects.select_for_update()
                .filter(owner_limit_query)
                .distinct()
                .count()
            )
            if owner_registered_count >= DOG_REGISTRATION_OWNER_MAX_PETS:
                _registration_error(
                    (
                        f"{owner_name} already has {DOG_REGISTRATION_OWNER_MAX_PETS} "
                        f"registered pets. A maximum of {DOG_REGISTRATION_OWNER_MAX_PETS} "
                        "pets is allowed per owner."
                    ),
                )
                return redirect('dogadoption_admin:register_dogs')

            dog = Dog.objects.create(
                date_registered=date_registered,
                name=name,
                species=species,
                sex=sex,
                age=age,
                neutering_status=neutering,
                color=color,
                owner_name=owner_name,
                owner_name_key=owner_name_key,
                owner_user=owner_user,
                owner_address=formatted_address,
                barangay=barangay,
            )
            for image_file in uploaded_images:
                DogImage.objects.create(dog=dog, image=image_file)

        cache.delete("registration_record_available_years")
        cache.delete("registration_record_active_barangays")
        image_suffix = f" with {len(uploaded_images)} photo(s)" if uploaded_images else ""
        _registration_success(f"Dog '{name}' registered successfully{image_suffix}!")
        return redirect('dogadoption_admin:register_dogs')

    return render(request, 'admin_registration/registration.html', {
        'barangay': barangay,
        'date': date
    })


def _get_cached_registration_barangays():
    barangay_list_parsed = cache.get("registration_record_active_barangays")
    if barangay_list_parsed is None:
        barangay_list_parsed = list(
            Barangay.objects.filter(is_active=True).values_list('name', flat=True)
        )
        cache.set("registration_record_active_barangays", barangay_list_parsed, 300)
    return barangay_list_parsed


def _get_cached_registration_years():
    available_years = cache.get("registration_record_available_years")
    if available_years is None:
        available_years = [
            d.year for d in Dog.objects.exclude(date_registered__isnull=True)
            .dates('date_registered', 'year', order='DESC')
        ]
        cache.set("registration_record_available_years", available_years, 300)
    return available_years


def _build_registration_record_queryset(selected_barangay, date_filter_type, filter_date, filter_month, filter_year):
    dogs = Dog.objects.all()
    if selected_barangay:
        dogs = dogs.filter(barangay__iexact=selected_barangay)

    dogs, date_filter_type, date_filter_label = _apply_registration_date_filter(
        dogs,
        date_filter_type,
        filter_date,
        filter_month,
        filter_year,
    )
    owner_group_key = Case(
        When(
            owner_user_id__isnull=False,
            then=Concat(
                Value("user:"),
                Cast("owner_user_id", output_field=CharField()),
                output_field=CharField(),
            ),
        ),
        When(
            Q(owner_name_key__isnull=False) & ~Q(owner_name_key=""),
            then=Concat(
                Value("name:"),
                Lower(Trim("owner_name_key")),
                output_field=CharField(),
            ),
        ),
        When(
            Q(owner_name__isnull=False) & ~Q(owner_name=""),
            then=Concat(
                Value("name:"),
                Lower(Trim("owner_name")),
                output_field=CharField(),
            ),
        ),
        default=Concat(
            Value("dog:"),
            Cast("id", output_field=CharField()),
            output_field=CharField(),
        ),
        output_field=CharField(),
    )
    dogs = dogs.annotate(owner_group_key=owner_group_key)
    owner_grouped_dogs = dogs
    owner_first_seen_date_subquery = (
        owner_grouped_dogs.filter(owner_group_key=OuterRef("owner_group_key"))
        .values("owner_group_key")
        .annotate(first_seen_date=Min("date_registered"))
        .values("first_seen_date")[:1]
    )
    dogs = dogs.annotate(
        owner_first_seen_date=Subquery(owner_first_seen_date_subquery)
    )
    owner_first_seen_id_subquery = (
        owner_grouped_dogs.filter(
            owner_group_key=OuterRef("owner_group_key"),
            date_registered=OuterRef("owner_first_seen_date"),
        )
        .order_by("id")
        .values("id")[:1]
    )
    dogs = (
        dogs.annotate(owner_first_seen_id=Subquery(owner_first_seen_id_subquery))
        .select_related("owner_user", "owner_user__profile")
        .only(
            "id",
            "date_registered",
            "name",
            "species",
            "sex",
            "age",
            "neutering_status",
            "color",
            "owner_name",
            "owner_name_key",
            "owner_address",
            "barangay",
            "owner_user_id",
            "owner_user__profile__profile_image",
        )
        .order_by("owner_first_seen_date", "owner_first_seen_id", "date_registered", "id")
    )
    return dogs, date_filter_type, date_filter_label


def _build_registration_owner_rank_lookup(dogs_qs):
    owner_rows = (
        dogs_qs.values("owner_group_key", "owner_first_seen_date", "owner_first_seen_id")
        .distinct()
        .order_by("owner_first_seen_date", "owner_first_seen_id", "owner_group_key")
    )
    return {
        row["owner_group_key"]: index
        for index, row in enumerate(owner_rows, start=1)
    }


def _attach_registration_owner_metadata(dogs):
    owner_profile_by_user_id = {}
    for dog in dogs:
        owner_user = getattr(dog, "owner_user", None)
        if not owner_user or not dog.owner_user_id:
            continue
        profile = _get_profile_or_none(owner_user)
        image_url = _safe_media_url(getattr(profile, "profile_image", None))
        if image_url and dog.owner_user_id not in owner_profile_by_user_id:
            owner_profile_by_user_id[dog.owner_user_id] = image_url

    names_without_user_profile = [
        dog.owner_name
        for dog in dogs
        if dog.owner_name and not dog.owner_user_id
    ]
    owner_profile_lookup = _build_owner_profile_lookup(names_without_user_profile)
    default_owner_profile_image_url = static("images/default-user-image.jpg")
    owner_keys_by_dog_id = {}

    for dog in dogs:
        normalized_owner = _normalize_person_name(dog.owner_name)
        matched_owner_profile = (
            owner_profile_lookup.get(normalized_owner, {})
            if not dog.owner_user_id
            else {}
        )
        matched_owner_user_id = matched_owner_profile.get("user_id")
        owner_key = getattr(dog, "owner_group_key", "") or _build_registration_record_owner_key(
            dog,
            matched_owner_user_id,
        )
        dog.owner_profile_image_url = (
            owner_profile_by_user_id.get(dog.owner_user_id)
            or matched_owner_profile.get("image_url", "")
            or default_owner_profile_image_url
        )
        dog.owner_profile_user_id = dog.owner_user_id or matched_owner_user_id
        if dog.owner_profile_user_id:
            dog.owner_profile_url = reverse(
                "dogadoption_admin:registration_owner_profile",
                args=[dog.owner_profile_user_id],
            )
        else:
            manual_params = {
                "owner_key": dog.owner_name_key or normalized_owner,
                "owner_name": dog.owner_name or "",
            }
            dog.owner_profile_url = (
                f"{reverse('dogadoption_admin:registration_owner_profile', args=[0])}"
                f"?{urlencode(manual_params)}"
            )
        dog.owner_initials = _owner_initials(dog.owner_name)
        owner_keys_by_dog_id[dog.id] = owner_key

    return owner_keys_by_dog_id


def _attach_registration_vaccination_metadata(dogs):
    today = timezone.localdate()
    candidate_owner_keys = set()
    candidate_pet_keys = set()
    dog_signature_by_id = {}

    for dog in dogs:
        owner_key = _normalize_person_name(
            getattr(dog, "owner_name_key", "") or getattr(dog, "owner_name", "")
        )
        pet_key = _normalize_person_name(getattr(dog, "name", ""))
        signature = (owner_key, pet_key)
        dog_signature_by_id[dog.id] = signature
        dog.has_vaccination_record = False
        dog.vaccination_expired = False
        dog.latest_vaccination_date = None
        dog.latest_vaccination_expiry_date = None

        if owner_key and pet_key:
            candidate_owner_keys.add(owner_key)
            candidate_pet_keys.add(pet_key)

    if not candidate_owner_keys or not candidate_pet_keys:
        return dogs

    matching_registrations = (
        DogRegistration.objects.annotate(
            owner_name_normalized=Lower(Trim("owner_name")),
            pet_name_normalized=Lower(Trim("name_of_pet")),
        )
        .filter(
            owner_name_normalized__in=candidate_owner_keys,
            pet_name_normalized__in=candidate_pet_keys,
        )
        .only("id", "date_registered")
    )

    registration_signature_by_id = {}
    registration_ids = []
    for registration in matching_registrations:
        signature = (
            getattr(registration, "owner_name_normalized", ""),
            getattr(registration, "pet_name_normalized", ""),
        )
        if not signature[0] or not signature[1]:
            continue
        registration_signature_by_id[registration.id] = signature
        registration_ids.append(registration.id)

    if not registration_ids:
        return dogs

    latest_vaccination_by_signature = {}
    vaccinations = (
        VaccinationRecord.objects.filter(registration_id__in=registration_ids)
        .only(
            "id",
            "registration_id",
            "date",
            "vaccine_expiry_date",
            "vaccination_expiry_date",
        )
        .order_by("-date", "-id")
    )

    for vaccination in vaccinations:
        signature = registration_signature_by_id.get(vaccination.registration_id)
        if signature and signature not in latest_vaccination_by_signature:
            latest_vaccination_by_signature[signature] = vaccination

    for dog in dogs:
        signature = dog_signature_by_id.get(dog.id)
        vaccination = latest_vaccination_by_signature.get(signature)
        if not vaccination:
            continue

        dog.has_vaccination_record = True
        dog.latest_vaccination_date = vaccination.date
        dog.latest_vaccination_expiry_date = (
            vaccination.vaccination_expiry_date or vaccination.vaccine_expiry_date
        )
        dog.vaccination_expired = bool(
            (vaccination.vaccine_expiry_date and vaccination.vaccine_expiry_date < today)
            or (
                vaccination.vaccination_expiry_date
                and vaccination.vaccination_expiry_date < today
            )
        )

    return dogs


def _paginate_registration_record_rows(request, dogs_qs):
    page_number = (request.GET.get("page") or "1").strip()
    paginator = Paginator(dogs_qs, 100)
    page_obj = paginator.get_page(page_number)
    paged_dogs = list(page_obj.object_list)

    return page_obj, paged_dogs


def _apply_registration_owner_row_display(dogs, owner_keys_by_dog_id, owner_rank_by_key):
    previous_owner_key = None
    for dog in dogs:
        owner_key = owner_keys_by_dog_id.get(dog.id, f"dog:{dog.id}")
        if owner_key == previous_owner_key:
            dog.owner_display_number = ""
            dog.show_owner_fields = False
        else:
            rank_key = getattr(dog, "owner_group_key", "") or owner_key
            dog.owner_display_number = owner_rank_by_key.get(rank_key, "")
            dog.show_owner_fields = True
        previous_owner_key = owner_key

    return dogs




# ---------------------------------------------------------------------------
# Register link 2/5: Registration List
# ---------------------------------------------------------------------------
@admin_required
def registration_record(request):
    """Render the grouped registration list with owner-level presentation data."""
    selected_barangay_raw = request.GET.get('barangay', '').strip()
    selected_barangay = _resolve_barangay_name(selected_barangay_raw) if selected_barangay_raw else ''
    date_filter_type = (request.GET.get('date_filter_type') or 'all').strip().lower()
    filter_date = (request.GET.get('filter_date') or '').strip()
    filter_month = (request.GET.get('filter_month') or '').strip()
    filter_year = (request.GET.get('filter_year') or '').strip()

    if date_filter_type not in {'all', 'day', 'month', 'year'}:
        date_filter_type = 'all'

    barangay_list_parsed = _get_cached_registration_barangays()
    dogs_qs, date_filter_type, date_filter_label = _build_registration_record_queryset(
        selected_barangay,
        date_filter_type,
        filter_date,
        filter_month,
        filter_year,
    )
    owner_rank_by_key = _build_registration_owner_rank_lookup(dogs_qs)
    page_obj, dogs = _paginate_registration_record_rows(request, dogs_qs)
    owner_keys_by_dog_id = _attach_registration_owner_metadata(dogs)
    dogs = _attach_registration_vaccination_metadata(dogs)
    dogs = _apply_registration_owner_row_display(
        dogs,
        owner_keys_by_dog_id,
        owner_rank_by_key,
    )
    available_years = _get_cached_registration_years()

    date_filter_params = _build_registration_filter_params(
        date_filter_type,
        filter_date,
        filter_month,
        filter_year,
    )

    date_filter_query = urlencode(date_filter_params)
    download_params = {}
    if selected_barangay:
        download_params['barangay'] = selected_barangay
    download_params.update(date_filter_params)
    download_query = urlencode(download_params)

    context = {
        'selected_barangay': selected_barangay,
        'dogs': dogs,
        'barangay_list_parsed': barangay_list_parsed,
        'registration_locator_points': _build_registration_locator_points(barangay_list_parsed),
        'date_filter_type': date_filter_type,
        'filter_date': filter_date,
        'filter_month': filter_month,
        'filter_year': filter_year,
        'date_filter_label': date_filter_label,
        'available_years': available_years,
        'date_filter_query': date_filter_query,
        'download_query': download_query,
        'page_obj': page_obj,
    }

    return render(request, 'admin_registration/registration_record.html', context)


@admin_required
def registration_owner_profile(request, user_id):
    """Show the profile preview used when drilling into a registration owner."""
    if int(user_id) <= 0:
        owner_key = _normalize_person_name(request.GET.get("owner_key"))
        owner_name = " ".join((request.GET.get("owner_name") or "").split()).strip()

        manual_owner_dogs_qs = Dog.objects.all()
        if owner_key:
            manual_owner_dogs_qs = manual_owner_dogs_qs.filter(owner_name_key=owner_key)
        elif owner_name:
            manual_owner_dogs_qs = manual_owner_dogs_qs.filter(owner_name__iexact=owner_name)
        else:
            raise Http404("Owner not found.")

        manual_owner_dogs_qs = (
            manual_owner_dogs_qs.prefetch_related(
                _dog_image_prefetch()
            )
            .only(
                "id",
                "name",
                "species",
                "sex",
                "age",
                "neutering_status",
                "color",
                "date_registered",
                "owner_name",
                "owner_address",
                "barangay",
            )
            .order_by("-date_registered", "-id")
        )
        manual_owner_dogs = list(manual_owner_dogs_qs)
        if not manual_owner_dogs:
            raise Http404("Owner not found.")

        resolved_owner_name = (
            owner_name
            or manual_owner_dogs[0].owner_name
            or "Manual Owner"
        )
        owner_name_parts = [part for part in resolved_owner_name.split() if part]
        manual_profile_user = User(
            username=owner_key or "manual-owner",
            first_name=owner_name_parts[0] if owner_name_parts else resolved_owner_name,
            last_name=" ".join(owner_name_parts[1:]) if len(owner_name_parts) > 1 else "",
        )
        manual_profile = Profile(
            user=manual_profile_user,
            middle_initial="",
            address=manual_owner_dogs[0].owner_address or "",
            age=None,
            consent_given=True,
        )
        manual_profile.profile_image = SimpleNamespace(
            url=static("images/default-user-image.jpg")
        )
        registered_dogs = _build_registered_dog_payloads(manual_owner_dogs)

        context = {
            "profile_user": manual_profile_user,
            "profile": manual_profile,
            "registered_dogs": registered_dogs,
            "registered_dogs_total": len(registered_dogs),
            "violation_summary": {
                "manual": 0,
                "claims": 0,
                "citations": 0,
                "legacy_total": 0,
                "total": 0,
            },
            "allow_image_preview": bool(request.user.is_staff),
        }
        return render(request, "admin_user/profile_preview.html", context)

    profile_user = get_object_or_404(
        User.objects.only("id", "username", "first_name", "last_name", "date_joined"),
        pk=user_id,
        is_staff=False,
    )
    profile, _ = Profile.objects.get_or_create(
        user=profile_user,
        defaults={"address": "", "age": 18, "consent_given": True},
    )

    registered_dogs_qs = (
        Dog.objects.filter(owner_user=profile_user)
        .prefetch_related(
            _dog_image_prefetch()
        )
        .only(
            "id",
            "name",
            "species",
            "sex",
            "age",
            "neutering_status",
            "color",
            "date_registered",
            "owner_address",
            "barangay",
        )
        .order_by("-date_registered", "-id")
    )
    registered_dogs = _build_registered_dog_payloads(registered_dogs_qs)

    violation_summary = (
        _admin_user_management_queryset()
        .filter(pk=profile_user.pk)
        .values(
            "claim_violation_count",
            "citation_violation_count",
            "calculated_violations",
            "managed_violation_count",
            "effective_violation_count",
        )
        .first()
        or {
            "claim_violation_count": 0,
            "citation_violation_count": 0,
            "calculated_violations": 0,
            "managed_violation_count": 0,
            "effective_violation_count": 0,
        }
    )

    context = {
        "profile_user": profile_user,
        "profile": profile,
        "registered_dogs": registered_dogs,
        "registered_dogs_total": len(registered_dogs),
        "violation_summary": {
            "manual": violation_summary.get("managed_violation_count", 0),
            "claims": violation_summary.get("claim_violation_count", 0),
            "citations": violation_summary.get("citation_violation_count", 0),
            "legacy_total": violation_summary.get("calculated_violations", 0),
            "total": violation_summary.get("effective_violation_count", 0),
        },
        "allow_image_preview": bool(request.user.is_staff),
    }
    return render(request, "admin_user/profile_preview.html", context)


@admin_required
def barangay_list_api(request):
    """Return active barangay names for registration autocomplete widgets."""
    query = " ".join((request.GET.get("q") or "").split()).strip()
    limit = _parse_positive_int(request.GET.get("limit"), default=200, max_value=200)
    cache_key = f"active_barangay_names:{query.casefold()}:{limit}"
    barangays = cache.get(cache_key)
    if barangays is None:
        barangay_qs = Barangay.objects.filter(is_active=True)
        if query:
            barangay_qs = barangay_qs.filter(name__icontains=query)
        barangays = list(barangay_qs.values_list('name', flat=True)[:limit])
        cache.set(cache_key, barangays, 300)
    return _cacheable_json_response({"barangays": barangays}, max_age=300)


@admin_required
def registration_user_search_api(request):
    """Search user accounts and registered owners to prefill registration details."""
    query = " ".join((request.GET.get("q") or "").split()).strip()
    limit = _parse_positive_int(request.GET.get("limit"))
    if len(query) < 2:
        return _cacheable_json_response({"results": []}, max_age=30)

    cache_key = f"registration_user_search:{query.casefold()}:{limit}"
    cached = cache.get(cache_key)
    if cached is not None:
        return _cacheable_json_response({"results": cached}, max_age=60)

    results = _build_registration_owner_search_results(query, limit)
    cache.set(cache_key, results, 60)
    return _cacheable_json_response({"results": results}, max_age=60)


@admin_required
def _get_registration_export_queryset(request):
    """Build the filtered registration export queryset once for all formats."""
    selected_barangay_raw = request.GET.get('barangay', None)
    selected_barangay = _resolve_barangay_name(selected_barangay_raw) if selected_barangay_raw else None
    date_filter_type = (request.GET.get('date_filter_type') or 'all').strip().lower()
    filter_date = (request.GET.get('filter_date') or '').strip()
    filter_month = (request.GET.get('filter_month') or '').strip()
    filter_year = (request.GET.get('filter_year') or '').strip()

    dogs = Dog.objects.only(
        "id",
        "date_registered",
        "name",
        "species",
        "sex",
        "age",
        "neutering_status",
        "color",
        "owner_name",
        "owner_address",
        "barangay",
    )

    if selected_barangay:
        dogs = dogs.filter(barangay__iexact=selected_barangay)

    dogs, _, _ = _apply_registration_date_filter(
        dogs,
        date_filter_type,
        filter_date,
        filter_month,
        filter_year,
    )

    return dogs.order_by("date_registered", "id"), (selected_barangay or "All Barangays")


def _build_registration_excel_response(dogs, selected_barangay_label):
    """Render the registration export as an Excel workbook."""
    Workbook, Alignment, Border, Font, PatternFill, Side = _get_openpyxl_exports()
    wb = Workbook()
    ws = wb.active
    ws.title = "Dog Registrations"
    ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
    ws.page_setup.paperSize = ws.PAPERSIZE_LETTER
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0

    thin = Side(style='thin', color='000000')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    header_fill = PatternFill(fill_type='solid', fgColor='D9D9D9')
    group_fill = PatternFill(fill_type='solid', fgColor='CFCFCF')

    ws.merge_cells('A1:J1')
    ws['A1'] = 'National Rabies Prevention and Control Program'
    ws.merge_cells('A2:J2')
    ws['A2'] = 'Rabies Free Visayas Project'
    ws.merge_cells('A3:J3')
    ws['A3'] = 'Dog Registry and Vaccination Record'
    ws.merge_cells('A5:J5')
    ws['A5'] = f'Name of Barangay: {selected_barangay_label}'

    for cell_ref, size, bold in [('A1', 11, False), ('A2', 12, True), ('A3', 14, True), ('A5', 11, False)]:
        cell = ws[cell_ref]
        cell.font = Font(size=size, bold=bold)
        cell.alignment = Alignment(horizontal='center' if cell_ref != 'A5' else 'left')

    ws.merge_cells('A7:B7')
    ws['A7'] = 'Registration'
    ws.merge_cells('C7:H7')
    ws['C7'] = 'Dog Profile'
    ws.merge_cells('I7:J7')
    ws['I7'] = 'Pet Owner'

    header_labels = [
        'No.', 'Date', 'Name', 'Species', 'Sex (M/F)',
        'Age (yrs)', 'Neutering (No./C/S)', 'Color',
        'Name', 'Complete Address'
    ]
    for col, label in enumerate(header_labels, start=1):
        ws.cell(row=8, column=col, value=label)

    for col in range(1, 11):
        group_cell = ws.cell(row=7, column=col)
        group_cell.fill = group_fill
        group_cell.font = Font(size=9, bold=True)
        group_cell.alignment = Alignment(horizontal='center', vertical='center')
        group_cell.border = border

        header_cell = ws.cell(row=8, column=col)
        header_cell.fill = header_fill
        header_cell.font = Font(size=8, bold=True)
        header_cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        header_cell.border = border

    start_data_row = 9
    for idx, dog in enumerate(dogs, start=1):
        row_idx = start_data_row + idx - 1
        row_values = [
            idx,
            dog.date_registered.strftime("%m-%d-%Y") if dog.date_registered else "",
            dog.name,
            dog.species,
            dog.sex,
            dog.age,
            dog.neutering_status,
            dog.color or "-",
            dog.owner_name,
            dog.owner_address,
        ]
        for col, value in enumerate(row_values, start=1):
            cell = ws.cell(row=row_idx, column=col, value=value)
            cell.font = Font(size=8)
            cell.border = border
            if col in (1, 2, 4, 5, 6, 7):
                cell.alignment = Alignment(horizontal='center', vertical='top', wrap_text=True)
            else:
                cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)

    last_data_row = start_data_row + max(len(dogs), 1) - 1
    legend_start_row = last_data_row + 2
    ws.cell(row=legend_start_row, column=7, value='Legend:').font = Font(size=8, bold=True)
    ws.cell(row=legend_start_row + 1, column=7, value='C - Castrated').font = Font(size=8)
    ws.cell(row=legend_start_row + 2, column=7, value='S - Spaying').font = Font(size=8)
    ws.cell(row=legend_start_row + 3, column=7, value='No - Not castrated nor spayed').font = Font(size=8)

    widths = [5, 11, 18, 10, 10, 11, 16, 13, 20, 36]
    for idx, width in enumerate(widths, start=1):
        ws.column_dimensions[chr(64 + idx)].width = width

    ws.print_title_rows = '1:8'

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    filename = f"Dog_Registrations_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    response['Content-Disposition'] = f'attachment; filename={filename}'
    wb.save(response)
    return response


def _build_registration_pdf_response(dogs, selected_barangay_label):
    """Render the registration export as a PDF report."""
    colors, landscape, letter, getSampleStyleSheet, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle = _get_reportlab_exports()
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(letter),
        leftMargin=18,
        rightMargin=18,
        topMargin=18,
        bottomMargin=20
    )
    styles = getSampleStyleSheet()
    title_style = styles['Heading4']
    title_style.alignment = 1
    title_style.fontSize = 11
    title_style.leading = 12

    small_center = styles['Normal'].clone('small_center')
    small_center.alignment = 1
    small_center.fontSize = 10

    small_left = styles['Normal'].clone('small_left')
    small_left.alignment = 0
    small_left.fontSize = 9

    elements = [
        Paragraph("National Rabies Prevention and Control Program", small_center),
        Paragraph("<b>Rabies Free Visayas Project</b>", small_center),
        Paragraph("<b>Dog Registry and Vaccination Record</b>", title_style),
        Spacer(1, 8),
        Paragraph(f"Name of Barangay: {selected_barangay_label}", small_left),
        Spacer(1, 6),
    ]

    pdf_cell = styles['Normal'].clone('pdf_cell')
    pdf_cell.fontSize = 7
    pdf_cell.leading = 8
    pdf_cell.wordWrap = 'CJK'

    data = [[
        'Registration', '', 'Dog Profile', '', '', '', '', '', 'Pet Owner', ''
    ], [
        'No.', 'Date', 'Name', 'Species', 'Sex (M/F)', 'Age (yrs)',
        'Neutering (No./C/S)', 'Color', 'Name', 'Complete Address'
    ]]

    for idx, dog in enumerate(dogs, start=1):
        data.append([
            idx,
            dog.date_registered.strftime("%m-%d-%Y") if dog.date_registered else "",
            dog.name,
            dog.species,
            dog.sex,
            dog.age,
            dog.neutering_status,
            dog.color or "-",
            Paragraph(dog.owner_name or "", pdf_cell),
            Paragraph(dog.owner_address or "", pdf_cell),
        ])

    col_widths = [28, 48, 76, 46, 42, 44, 74, 50, 80, 242]
    table = Table(data, colWidths=col_widths, repeatRows=0)
    table.setStyle(TableStyle([
        ('SPAN', (0, 0), (1, 0)),
        ('SPAN', (2, 0), (7, 0)),
        ('SPAN', (8, 0), (9, 0)),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#d0d0d0')),
        ('BACKGROUND', (0, 1), (-1, 1), colors.HexColor('#e6e6e6')),
        ('FONTNAME', (0, 0), (-1, 1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 7.5),
        ('ALIGN', (0, 0), (-1, 1), 'CENTER'),
        ('ALIGN', (0, 2), (1, -1), 'CENTER'),
        ('ALIGN', (3, 2), (7, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('WORDWRAP', (8, 2), (9, -1), 'CJK'),
        ('GRID', (0, 0), (-1, -1), 0.7, colors.black),
        ('TOPPADDING', (0, 0), (-1, -1), 2),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ('LEFTPADDING', (0, 0), (-1, -1), 3),
        ('RIGHTPADDING', (0, 0), (-1, -1), 3),
    ]))

    elements.append(table)
    elements.append(Spacer(1, 8))
    elements.append(Paragraph("Legend:", small_left))
    elements.append(Paragraph("C - Castrated", small_left))
    elements.append(Paragraph("S - Spaying", small_left))
    elements.append(Paragraph("No - Not castrated nor spayed", small_left))

    doc.build(elements)
    buffer.seek(0)

    response = HttpResponse(buffer, content_type='application/pdf')
    filename = f"Dog_Registrations_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    response['Content-Disposition'] = f'attachment; filename={filename}'
    return response


@admin_required
def download_registration(request, file_type):
    """Export the filtered registration list as Excel or PDF."""
    dogs, selected_barangay_label = _get_registration_export_queryset(request)

    if file_type == 'excel':
        return _build_registration_excel_response(dogs, selected_barangay_label)
    if file_type == 'pdf':
        return _build_registration_pdf_response(dogs, selected_barangay_label)

    return HttpResponse("Invalid file type.", status=400)


# ---------------------------------------------------------------------------
# Register link 3/5: Vaccination
# ---------------------------------------------------------------------------
@admin_required
def med_record(request, registration_id):
    """Maintain vaccination and deworming records for one registration."""
    registration = DogRegistration.objects.get(id=registration_id)

    vaccinations = VaccinationRecord.objects.filter(
        registration=registration
    ).order_by('-date')

    dewormings = DewormingTreatmentRecord.objects.filter(
        registration=registration
    ).order_by('-date')

    current_address = (registration.address or "").strip()
    current_street = ""
    current_barangay = ""
    if current_address:
        parts = [p.strip() for p in current_address.split(",") if p.strip()]
        if len(parts) >= 3 and parts[-2].lower() == "bayawan city" and parts[-1].lower() == "negros oriental":
            current_barangay = parts[-3]
            current_street = ", ".join(parts[:-3])
        else:
            matched_barangay = _resolve_barangay_name(current_address)
            current_barangay = matched_barangay
            if matched_barangay:
                current_street = current_address.replace(matched_barangay, "").strip(" ,")

    cert_settings = CertificateSettings.objects.first()
    vaccination_defaults = {
        "vac_date": cert_settings.default_vac_date.isoformat() if cert_settings and cert_settings.default_vac_date else "",
        "vaccine_name": cert_settings.default_vaccine_name if cert_settings else "",
        "manufacturer_lot_no": cert_settings.default_manufacturer_lot_no if cert_settings else "",
        "vaccine_expiry_date": cert_settings.default_vaccine_expiry_date.isoformat() if cert_settings and cert_settings.default_vaccine_expiry_date else "",
    }

    if request.method == "POST":
        record_type = request.POST.get("record_type")

        if record_type == "update_address":
            barangay_input = request.POST.get("barangay", "")
            street_address = (request.POST.get("street_address") or "").strip()
            barangay = _resolve_barangay_name(barangay_input)

            if not barangay:
                messages.error(request, "Please select a valid barangay from the suggestions.")
                return redirect('dogadoption_admin:med_records', registration_id=registration.id)

            registration.address = (
                f"{street_address}, {barangay}, Bayawan City, Negros Oriental"
                if street_address else
                f"{barangay}, Bayawan City, Negros Oriental"
            )
            registration.save(update_fields=["address"])
            messages.success(request, "Owner address updated.")
            return redirect('dogadoption_admin:med_records', registration_id=registration.id)

        if record_type == "vaccination":
            (
                vac_date,
                vaccine_name,
                manufacturer_lot_no,
                vaccine_expiry_date,
                vaccination_expiry_date,
            ) = _get_vaccination_post_values(request)
            cert_settings = _create_vaccination_and_update_defaults(
                registration,
                cert_settings,
                vac_date,
                vaccine_name,
                manufacturer_lot_no,
                vaccine_expiry_date,
                vaccination_expiry_date,
            )

        elif record_type == "deworming":
            DewormingTreatmentRecord.objects.create(
                registration=registration,
                date=request.POST.get("dew_date"),
                medicine_given=request.POST.get("medicine_given"),
                medicine_expiry_date=(request.POST.get("medicine_expiry_date") or "").strip() or None,
                route=request.POST.get("route"),
                frequency=request.POST.get("frequency"),
                veterinarian=request.POST.get("dew_veterinarian"),
            )

        elif record_type == "all":
            (
                vac_date,
                vaccine_name,
                manufacturer_lot_no,
                vaccine_expiry_date,
                vaccination_expiry_date,
            ) = _get_vaccination_post_values(request)
            cert_settings = _create_vaccination_and_update_defaults(
                registration,
                cert_settings,
                vac_date,
                vaccine_name,
                manufacturer_lot_no,
                vaccine_expiry_date,
                vaccination_expiry_date,
            )

            DewormingTreatmentRecord.objects.create(
                registration=registration,
                date=request.POST.get("dew_date"),
                medicine_given=request.POST.get("medicine_given"),
                medicine_expiry_date=(request.POST.get("medicine_expiry_date") or "").strip() or None,
                route=request.POST.get("route"),
                frequency=request.POST.get("frequency"),
                veterinarian=request.POST.get("dew_veterinarian"),
            )

        sync_expiry_notifications()
        cache.delete(ADMIN_NOTIFICATIONS_CACHE_KEY)
        invalidate_user_notification_content()
        return redirect('dogadoption_admin:med_records', registration_id=registration.id)

    context = {
        "registration": registration,
        "vaccinations": vaccinations,
        "dewormings": dewormings,
        "current_street": current_street,
        "current_barangay": current_barangay,
        "vaccination_defaults": vaccination_defaults,
    }

    return render(request, "admin_registration/med_record.html", context)

@admin_required
def dog_certificate(request):
    """Create the base vaccination certificate registration before medical entry."""
    settings = CertificateSettings.objects.first()

    if request.method == "POST":
        reg_no = (request.POST.get("reg_no") or "").strip()
        series_prefix = _normalize_certificate_series(reg_no)
        breed = (request.POST.get("breed") or "").strip()
        dob_input = (request.POST.get("dob") or "").strip()
        dob_value = parse_date(dob_input) if dob_input else None
        barangay_input = request.POST.get("barangay", "")
        barangay = _resolve_barangay_name(barangay_input)
        address_line = (request.POST.get("address") or "").strip()
        owner_first_name = (request.POST.get("owner_first_name") or "").strip()
        owner_last_name = (request.POST.get("owner_last_name") or "").strip()
        owner_user_id = (request.POST.get("owner_user_id") or "").strip()
        owner_name = _build_owner_full_name(
            owner_first_name,
            owner_last_name,
            (request.POST.get("owner_name") or "").strip(),
        )
        owner_name, owner_name_key, _owner_user = _resolve_registration_owner_identity(
            owner_first_name,
            owner_last_name,
            owner_user_id=owner_user_id,
        )
        status = (request.POST.get("status") or "").strip()

        if not series_prefix:
            messages.error(request, "Please enter a valid registration series.")
            return render(request, 'admin_registration/dog_certificate.html', {'settings': settings})

        if not breed:
            messages.error(request, "Breed is required.")
            return render(request, 'admin_registration/dog_certificate.html', {'settings': settings})

        if dob_input and not dob_value:
            messages.error(request, "Please enter a valid Date of Birth.")
            return render(request, 'admin_registration/dog_certificate.html', {'settings': settings})

        if status not in {"Castrated", "Spayed", "Intact"}:
            messages.error(request, "Please select a valid status.")
            return render(request, 'admin_registration/dog_certificate.html', {'settings': settings})

        if not owner_name or not owner_name_key:
            messages.error(request, "Owner First and Last Name are required.")
            return render(request, 'admin_registration/dog_certificate.html', {'settings': settings})

        if (owner_first_name or owner_last_name) and (not owner_first_name or not owner_last_name):
            messages.error(request, "Please provide both owner first name and last name.")
            return render(request, 'admin_registration/dog_certificate.html', {'settings': settings})

        if not barangay:
            messages.error(request, "Please select a valid barangay from the suggestions.")
            return render(request, 'admin_registration/dog_certificate.html', {'settings': settings})

        contact_no_input = (request.POST.get("contact_no") or "").strip()
        contact_no_digits = re.sub(r"\D", "", contact_no_input)

        # Accept common PH mobile formats, including spaced numbers like 0912 345 6789.
        canonical_local = ""
        if len(contact_no_digits) == 11 and contact_no_digits.startswith("09"):
            canonical_local = contact_no_digits
        elif len(contact_no_digits) == 10 and contact_no_digits.startswith("9"):
            canonical_local = f"0{contact_no_digits}"
        elif len(contact_no_digits) == 12 and contact_no_digits.startswith("639"):
            canonical_local = f"0{contact_no_digits[2:]}"

        if not re.match(r"^09\d{9}$", canonical_local):
            messages.error(request, "Use a valid PH mobile number: 09XXXXXXXXX (spaces are allowed).")
            return render(request, 'admin_registration/dog_certificate.html', {'settings': settings})

        contact_no = f"+63{canonical_local[1:]}"

        with transaction.atomic():
            settings = (
                CertificateSettings.objects.select_for_update()
                .order_by("id")
                .first()
            )
            if settings:
                if settings.reg_no != series_prefix:
                    settings.reg_no = series_prefix
                    settings.save(update_fields=["reg_no"])
            else:
                settings = CertificateSettings.objects.create(reg_no=series_prefix)

            registration = DogRegistration.objects.create(
                # Keep address format standardized: "<Barangay>, Bayawan City, Negros Oriental"
                reg_no=_build_certificate_registration_number(settings.reg_no),
                name_of_pet=request.POST.get('name_of_pet'),
                breed=breed,
                dob=dob_value,
                color_markings=request.POST.get('color_markings'),
                sex=request.POST.get('sex'),
                status=status,
                owner_name=owner_name,
                address=f"{address_line}, {barangay}, Bayawan City, Negros Oriental" if address_line else f"{barangay}, Bayawan City, Negros Oriental",
                contact_no=contact_no,
            )

        #  Redirect to medical record with dog ID
        return redirect('dogadoption_admin:med_records', registration_id=registration.id)

    return render(request, 'admin_registration/dog_certificate.html', {'settings': settings})

# ---------------------------------------------------------------------------
# Register link 4/5: Vaccination List
# ---------------------------------------------------------------------------
@admin_required
def certificate_print(request, pk):
    """Render a printable certificate for one registration."""
    registration = get_object_or_404(DogRegistration, pk=pk)

    vaccinations = VaccinationRecord.objects.filter(
        registration=registration
    ).order_by('-date')

    dewormings = DewormingTreatmentRecord.objects.filter(
        registration=registration
    ).order_by('-date')

    context = {
        'certificate': _build_certificate_payload(
            registration,
            vaccinations=vaccinations,
            dewormings=dewormings,
            vac_limit=3,
            vac_min_rows=3,
            dew_limit=3,
            dew_min_rows=3,
        ),
    }

    return render(request, 'admin_registration/certificate_print.html', context)

@admin_required
def certificate_list(request):
    """Show issued vaccination certificates and expiry tracking by barangay."""
    today = timezone.localdate()
    barangay_names = list(
        Barangay.objects.filter(is_active=True)
        .order_by('sort_order', 'name')
        .values_list('name', flat=True)
    )
    selected_query = " ".join(
        (
            request.GET.get("q")
            or request.GET.get("barangay")
            or ""
        ).split()
    ).strip()
    selected_sort = " ".join((request.GET.get("sort") or "reg_no").split()).strip().lower()
    if selected_sort not in {"reg_no", "owner_user"}:
        selected_sort = "reg_no"

    medical_records_qs = (
        VaccinationRecord.objects.select_related('registration')
        .filter(registration__isnull=False)
    )

    if selected_query:
        medical_records_qs = medical_records_qs.filter(
            Q(registration__reg_no__icontains=selected_query)
            | Q(registration__name_of_pet__icontains=selected_query)
            | Q(registration__owner_name__icontains=selected_query)
            | Q(registration__address__icontains=selected_query)
            | Q(vaccine_name__icontains=selected_query)
        )

    combined_rows = []
    for record in medical_records_qs:
        reg = record.registration
        combined_rows.append({
            'registration_id': reg.id,
            'reg_no': reg.reg_no,
            'pet_name': reg.name_of_pet,
            'owner_name': reg.owner_name,
            'barangay': _extract_barangay_from_address(reg.address) or '-',
            'address': reg.address,
            'date_issued': reg.date_registered,
            'vaccination_date': record.date,
            'vaccine_name': record.vaccine_name,
            'vaccine_expiry_date': record.vaccine_expiry_date,
            'dog_vaccination_expiry_date': record.vaccination_expiry_date,
            'is_expired': (
                record.vaccine_expiry_date < today or
                record.vaccination_expiry_date < today
            ),
        })

    _attach_certificate_owner_metadata(combined_rows)
    combined_rows.sort(
        key=lambda row: (
            _certificate_reg_no_sort_key(row.get("reg_no")),
            row.get("vaccination_date") or datetime.min.date(),
            row.get("registration_id") or 0,
        ),
        reverse=True,
    )

    owner_groups = _build_certificate_owner_groups(combined_rows)
    page_items = owner_groups if selected_sort == "owner_user" else combined_rows
    page_obj = Paginator(page_items, 10).get_page(request.GET.get('page'))

    expired_vaccinations = (
        VaccinationRecord.objects.select_related('registration')
        .filter(
            Q(vaccine_expiry_date__lt=today) |
            Q(vaccination_expiry_date__lt=today)
        )
        .order_by('vaccine_expiry_date', 'vaccination_expiry_date')
    )

    tracker_map = {name: [] for name in barangay_names}

    for row in expired_vaccinations:
        barangay_name = _extract_barangay_from_address(row.registration.address)
        if barangay_name not in tracker_map:
            continue

        tracker_map[barangay_name].append({
            'reg_no': row.registration.reg_no,
            'owner_name': row.registration.owner_name,
            'vaccine_name': row.vaccine_name,
            'vaccine_expiry_date': row.vaccine_expiry_date,
            'dog_vaccination_expiry_date': row.vaccination_expiry_date,
        })

    barangay_expiry_tracker = [
        {
            'barangay': name,
            'expired_count': len(tracker_map[name]),
            'expired_items': tracker_map[name],
        }
        for name in barangay_names
    ]

    return render(request, 'admin_registration/certificate_list.html', {
        'page_obj': page_obj,
        'selected_query': selected_query,
        'selected_sort': selected_sort,
        'sort_options': (
            ("reg_no", "Reg. No."),
            ("owner_user", "Owner / User"),
        ),
        'is_grouped_sort': selected_sort == "owner_user",
        'total_vaccination_records': len(combined_rows),
        'total_owner_groups': len(owner_groups),
        'barangay_names': barangay_names,
        'barangay_expiry_tracker': barangay_expiry_tracker,
    })


@admin_required
def export_certificates_pdf(request):
    """Export the certificate list as a compact PDF table."""
    _, _, _, _, _, SimpleDocTemplate, _, Table, _ = _get_reportlab_exports()
    certificates = DogRegistration.objects.all().order_by('-date_registered')

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="certificates.pdf"'

    doc = SimpleDocTemplate(response)
    data = [["Reg No", "Pet Name", "Owner", "Date Issued"]]

    for cert in certificates:
        data.append([
            cert.reg_no,
            cert.name_of_pet,
            cert.owner_name,
            cert.date_registered.strftime("%b %d, %Y")
        ])

    table = Table(data)
    doc.build([table])
    return response

@admin_required
def export_certificates_word(request):
    """Export the certificate list as a Word document."""
    Document = _get_python_docx_document()
    certificates = DogRegistration.objects.all().order_by('-date_registered')

    document = Document()
    document.add_heading('Vaccination Certificates', level=1)

    table = document.add_table(rows=1, cols=4)
    headers = ["Reg No", "Pet Name", "Owner", "Date Issued"]

    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    for cert in certificates:
        row_cells = table.add_row().cells
        row_cells[0].text = cert.reg_no
        row_cells[1].text = cert.name_of_pet
        row_cells[2].text = cert.owner_name
        row_cells[3].text = cert.date_registered.strftime("%b %d, %Y")

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="certificates.docx"'
    document.save(response)
    return response

@admin_required
def export_certificates_excel(request):
    """Export the certificate list as an Excel workbook."""
    Workbook, _, _, _, _, _ = _get_openpyxl_exports()
    certificates = DogRegistration.objects.all().order_by('-date_registered')

    wb = Workbook()
    ws = wb.active
    ws.title = "Certificates"

    ws.append(["Reg No", "Pet Name", "Owner", "Date Issued"])

    for cert in certificates:
        ws.append([
            cert.reg_no,
            cert.name_of_pet,
            cert.owner_name,
            cert.date_registered.strftime("%b %d, %Y")
        ])

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="certificates.xlsx"'
    wb.save(response)
    return response


@admin_required
def bulk_certificate_print(request):
    """Render multiple printable certificates from the selected list rows."""
    if request.method == "POST":
        selected_ids = [int(pk) for pk in request.POST.getlist("selected_ids") if str(pk).isdigit()]

        if not selected_ids:
            return redirect("dogadoption_admin:certificate_list")

        # ONLY fetch the selected certificates
        registrations = (
            DogRegistration.objects.filter(id__in=selected_ids)
            .prefetch_related(
                Prefetch(
                    "vaccinations",
                    queryset=VaccinationRecord.objects.order_by("-date"),
                    to_attr="vaccination_records_sorted",
                ),
                Prefetch(
                    "dewormings",
                    queryset=DewormingTreatmentRecord.objects.order_by("-date"),
                    to_attr="deworming_records_sorted",
                ),
            )
            .order_by('id')
        )

        certificates = [
            _build_certificate_payload(
                registration,
                vaccinations=getattr(registration, "vaccination_records_sorted", []),
                dewormings=getattr(registration, "deworming_records_sorted", []),
                vac_limit=3,
                vac_min_rows=3,
                dew_limit=3,
                dew_min_rows=3,
            )
            for registration in registrations
        ]

        return render(
            request,
            "admin_registration/certificate_print.html",
            {"certificates": certificates}
        )

    return redirect("dogadoption_admin:certificate_list")


# ---------------------------------------------------------------------------
# Register link 5/5: Citation
# ---------------------------------------------------------------------------
def citation_create(request):
    """Create a citation and show the latest citation activity summary."""
    form = CitationForm(request.POST or None)
    latest_citation = Citation.objects.order_by('-id').first()
    today = timezone.localdate()
    citations_qs = Citation.objects.select_related('owner', 'penalty') \
        .prefetch_related('penalties', 'penalties__section') \
        .order_by('-id')[:10]
    citation_rows = []
    for citation in citations_qs:
        penalties = list(citation.penalties.all())
        if not penalties and citation.penalty_id:
            penalties = [citation.penalty]
        violations = ", ".join([p.title for p in penalties]) if penalties else "-"
        total_fees = sum([p.amount for p in penalties]) if penalties else 0
        citation_name = " ".join(
            part for part in [citation.owner_first_name, citation.owner_last_name] if part
        ).strip()
        if not citation_name and citation.owner_id:
            citation_name = citation.owner.get_full_name().strip() or citation.owner.username
        if not citation_name:
            citation_name = "Unknown Owner"
        citation_rows.append({
            "citation": citation,
            "display_name": citation_name,
            "violations": violations,
            "total_fees": total_fees,
        })
    penalties = Penalty.objects.filter(active=True).select_related('section').order_by('section__number', 'number')

    owner_search_data = []
    owner_rows = User.objects.filter(is_staff=False).values(
        "id",
        "username",
        "first_name",
        "last_name",
        "profile__address",
    ).order_by("username")
    for row in owner_rows:
        owner_search_data.append({
            "id": row["id"],
            "username": row["username"] or "",
            "first_name": row["first_name"] or "",
            "last_name": row["last_name"] or "",
            "barangay": _extract_barangay_from_address(row.get("profile__address") or ""),
        })

    today_claim_requests = []
    seen_today_request_users = set()
    today_request_qs = (
        PostRequest.objects.filter(
            request_type="claim",
            status="accepted",
            scheduled_appointment_date=today,
            user__is_staff=False,
        )
        .select_related("user", "user__profile")
        .order_by("user__username", "id")
    )
    for req in today_request_qs:
        if req.user_id in seen_today_request_users:
            continue
        seen_today_request_users.add(req.user_id)
        owner_profile = getattr(req.user, "profile", None)
        today_claim_requests.append({
            "user_id": req.user_id,
            "username": req.user.username or "",
            "first_name": req.user.first_name or "",
            "last_name": req.user.last_name or "",
            "barangay": _extract_barangay_from_address(getattr(owner_profile, "address", "") or ""),
        })

    if request.method == 'POST' and form.is_valid():
        selected_ids = request.POST.getlist('penalties')
        selected_penalties = list(Penalty.objects.filter(id__in=selected_ids, active=True).order_by('section__number', 'number'))

        if not selected_penalties:
            messages.error(request, 'Please select at least one violation.')
        else:
            citation = form.save(commit=False)
            owner = form.cleaned_data.get("owner")
            citation.owner = owner

            if owner and not citation.owner_barangay:
                owner_address = getattr(getattr(owner, "profile", None), "address", "") or ""
                citation.owner_barangay = _extract_barangay_from_address(owner_address)

            # Keep backward compatibility with existing single-penalty references.
            citation.penalty = selected_penalties[0]
            citation.save()
            citation.penalties.set(selected_penalties)
            return redirect('dogadoption_admin:citation_print', citation.pk)

    return render(request, 'admin_registration/citation_form.html', {
        'form': form,
        'latest_citation': latest_citation,
        'citation_rows': citation_rows,
        'penalties': penalties,
        'owner_search_data': owner_search_data,
        'today_claim_requests': today_claim_requests,
        'today_claim_date': today,
        'selected_penalty_ids': [int(x) for x in request.POST.getlist('penalties') if str(x).isdigit()] if request.method == 'POST' else [],
    })

def citation_print(request, pk):
    """Render the printable citation view for a single citation record."""
    citation = get_object_or_404(Citation, pk=pk)
    selected_penalties = list(citation.penalties.all().select_related('section').order_by('section__number', 'number'))
    if not selected_penalties and citation.penalty_id:
        selected_penalties = [citation.penalty]

    owner_name = " ".join(
        part for part in [citation.owner_first_name, citation.owner_last_name] if part
    ).strip() or "Unknown Owner"
    owner_address = "-"
    owner_barangay = citation.owner_barangay or "-"
    if citation.owner_id:
        try:
            profile_address = citation.owner.profile.address or "-"
        except Exception:
            profile_address = "-"
        owner_address = profile_address

        if owner_name == "Unknown Owner":
            owner_name = citation.owner.get_full_name().strip() or citation.owner.username

        if not citation.owner_barangay:
            extracted_barangay = _extract_barangay_from_address(owner_address)
            if extracted_barangay:
                owner_barangay = extracted_barangay
            elif owner_address and owner_address != "-":
                owner_barangay = owner_address.split(",")[0].strip() or "-"

    total_amount = sum((p.amount for p in selected_penalties), Decimal("0.00"))
    receipt_seed = f"{citation.id}|{citation.owner_id or 0}|{citation.date_issued.isoformat()}"
    receipt_hash = hashlib.sha256(receipt_seed.encode("utf-8")).hexdigest()[:10].upper()
    receipt_no = f"CIT-{citation.id:06d}-{receipt_hash}"

    return render(request, 'admin_registration/citation_print.html', {
        'citation': citation,
        'selected_penalties': selected_penalties,
        'owner_name': owner_name,
        'owner_address': owner_address,
        'owner_barangay': owner_barangay,
        'total_amount': total_amount,
        'receipt_no': receipt_no,
    })

def penalty_manager(request):
    """Manage penalty sections and penalty items used by citations."""
    editing_penalty = None
    edit_penalty_id = request.GET.get('edit_penalty')
    if str(edit_penalty_id).isdigit():
        editing_penalty = get_object_or_404(
            Penalty.objects.select_related('section'),
            pk=int(edit_penalty_id),
        )

    s_form = SectionForm()
    p_form = PenaltyForm(instance=editing_penalty) if editing_penalty else PenaltyForm()

    if request.method == 'POST':
        if 'add_section' in request.POST:
            s_form = SectionForm(request.POST)
            if s_form.is_valid():
                s_form.save()
                messages.success(request, "Section added.")
                return redirect('dogadoption_admin:penalty_manage')

        elif 'add_penalty' in request.POST:
            p_form = PenaltyForm(request.POST)
            if p_form.is_valid():
                p_form.save()
                messages.success(request, "Penalty added.")
                return redirect('dogadoption_admin:penalty_manage')

        elif 'update_penalty' in request.POST:
            penalty_id = request.POST.get('penalty_id')
            editing_penalty = get_object_or_404(Penalty, pk=penalty_id)
            p_form = PenaltyForm(request.POST, instance=editing_penalty)
            if p_form.is_valid():
                p_form.save()
                messages.success(request, "Penalty updated.")
                return redirect('dogadoption_admin:penalty_manage')

        elif 'delete_penalty' in request.POST:
            penalty_id = request.POST.get('penalty_id')
            penalty = get_object_or_404(Penalty, pk=penalty_id)
            penalty.delete()
            messages.success(request, "Penalty deleted.")
            return redirect('dogadoption_admin:penalty_manage')

    sections = list(
        PenaltySection.objects.prefetch_related('penalties').order_by('number')
    )
    total_penalties = 0
    active_penalties = 0
    for section in sections:
        section_penalties = list(section.penalties.all().order_by('number'))
        section.penalties_list = section_penalties
        total_penalties += len(section_penalties)
        active_penalties += sum(1 for penalty in section_penalties if penalty.active)

    return render(request, 'admin_registration/penalty_manage.html', {
        'sections': sections,
        's_form': s_form,
        'p_form': p_form,
        'editing_penalty': editing_penalty,
        'section_count': len(sections),
        'penalty_count': total_penalties,
        'active_penalty_count': active_penalties,
        'inactive_penalty_count': max(total_penalties - active_penalties, 0),
    })


# ---------------------------------------------------------------------------
# User Post Requests (admin approval gate)
# ---------------------------------------------------------------------------

def _build_user_post_items(adoption_qs, missing_qs):
    """Merge adoption and missing post querysets into a single sorted list."""
    items = []
    for post in adoption_qs:
        first_image = post.images.first()
        items.append({
            "id": post.id,
            "post_type": "adoption",
            "post_type_label": "Adoption",
            "dog_name": post.dog_name,
            "breed": post.display_breed,
            "age": post.age,
            "description": post.description,
            "location": post.location,
            "status": post.status,
            "image_url": first_image.image.url if first_image else None,
            "owner_username": post.owner.username,
            "owner_full_name": post.owner.get_full_name() or post.owner.username,
            "created_at": post.created_at,
        })
    for post in missing_qs:
        items.append({
            "id": post.id,
            "post_type": "missing",
            "post_type_label": "Missing",
            "dog_name": post.dog_name,
            "age": post.age,
            "description": post.description,
            "location": post.location,
            "status": post.status,
            "image_url": post.image.url if post.image else None,
            "owner_username": post.owner.username,
            "owner_full_name": post.owner.get_full_name() or post.owner.username,
            "created_at": post.created_at,
        })
    items.sort(key=lambda x: x["created_at"], reverse=True)
    return items


@admin_required
def user_post_requests(request):
    """Admin page listing user-submitted adoption/missing posts for review."""
    tab = request.GET.get("tab", "pending")
    if tab not in ("pending", "accepted", "declined"):
        tab = "pending"

    adoption_filters = {
        "pending": {"status": "pending_review"},
        "accepted": {"status": "available"},
        "declined": {"status": "declined"},
    }
    missing_filters = {
        "pending": {"status": "pending_review"},
        "accepted": {"status": "missing"},
        "declined": {"status": "declined"},
    }

    adoption_qs = (
        UserAdoptionPost.objects
        .filter(**adoption_filters[tab])
        .select_related("owner")
        .prefetch_related("images")
        .order_by("-created_at")
    )
    missing_qs = (
        MissingDogPost.objects
        .filter(**missing_filters[tab])
        .select_related("owner")
        .order_by("-created_at")
    )

    items = _build_user_post_items(adoption_qs, missing_qs)

    page_number = request.GET.get("page", 1)
    paginator = Paginator(items, 12)
    page_obj = paginator.get_page(page_number)

    pending_count = (
        UserAdoptionPost.objects.filter(status="pending_review").count()
        + MissingDogPost.objects.filter(status="pending_review").count()
    )
    accepted_count = (
        UserAdoptionPost.objects.filter(status="available").count()
        + MissingDogPost.objects.filter(status="missing").count()
    )
    declined_count = (
        UserAdoptionPost.objects.filter(status="declined").count()
        + MissingDogPost.objects.filter(status="declined").count()
    )

    return render(request, "admin_home/user_post_requests.html", {
        "tab": tab,
        "page_obj": page_obj,
        "pending_count": pending_count,
        "accepted_count": accepted_count,
        "declined_count": declined_count,
    })


@require_POST
@admin_required
def user_post_request_action(request, post_type, post_id, action):
    """Handle accept / decline / delete actions on user-submitted posts."""
    if post_type == "adoption":
        post = get_object_or_404(UserAdoptionPost, pk=post_id)
        accept_status = "available"
    elif post_type == "missing":
        post = get_object_or_404(MissingDogPost, pk=post_id)
        accept_status = "missing"
    else:
        raise Http404

    if action == "accept":
        post.status = accept_status
        post.save(update_fields=["status"])
        messages.success(request, f"Post by {post.owner.username} has been accepted.")
    elif action == "decline":
        post.status = "declined"
        post.save(update_fields=["status"])
        messages.success(request, f"Post by {post.owner.username} has been declined.")
    elif action == "delete":
        owner_name = post.owner.username
        post.delete()
        messages.success(request, f"Post by {owner_name} has been deleted.")
    else:
        raise Http404

    return redirect(reverse("dogadoption_admin:user_post_requests"))
